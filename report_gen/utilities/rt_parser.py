# Risk & Vulnerability Assessment Reporting Engine

# Copyright 2022 The Risk & Vulnerability Reporting Engine Contributors, All Rights Reserved.
# (see Contributors.txt for a full list of Contributors)

# SPDX-License-Identifier: BSD-3-Clause

# Please see additional acknowledgments (including references to third party source code, object code, documentation and other files) in the license.txt file or contact permission@sei.cmu.edu for full terms.

# Created, in part, with funding and support from the United States Government. (see Acknowledgments file).

# DM22-1011

from docx import Document
from docx.shared import Inches
from html.parser import HTMLParser


class RichTextParser(HTMLParser):
    """RichTextParser

    The function will parse the HTML like string from the Rich Text Widget used
    on the front end and based via the serializer

    Args:
              paragraph (Docx paragraph): The paragraph in the document to insert the rich text before.
              debug (bool, optional): Flag for output debug messages. Defaults to False.
              text_style (str, optional): The style of the paragraphs that are created from the rich text.. Defaults to 'Normal'.

    Instantiate the class like

        parser = RichTextParser(paragraph)

     where paragraph is the location of the tag to be replaced in the current
     python-docx object.  There is also a debug flag available that prints
     information about the state of the parsed rich text being created in
     emit_docx

     Once instanciated, the object can be reused multiple times.  So you would
     call it like:

        parser.feed("<p><b>Boldly</b> go...</p>")
        parser.emit_docx()
        parser.reset_parser()

     internally, the code takes the result of the html parser and creates tuples
     marking the start and end of tags and then the actual text that will go in
     the python-docx runs.

     Note: there are 2 types of actions once the processing begins and the are
     marked with `emit' and `adj'.

       `emit' will be actual interactions with python-docx.
       `adj' adjusts the state of style or the depth/type of a list

    NB: ordered list do not reset currently do to the complexity of how word
    resets numbering. trust me, you don't want to know"""

    style_eles = ["b", "i", "u"]
    list_types = {"ul": "List Bullet 2", "ol": "List Number 2"}

    def __init__(self, paragraph, debug=False, text_style='Normal'):
        """Constructor for the rich text parser.

        Args:
            paragraph (Docx Paragraph or pptx TextFrame): Where the richtext will be placed. The object type will differ from docx to pptx.
            debug (bool, optional): Wether debug statements will be printed during parsing. Defaults to False.
            text_style (str, optional): Styling of the text. This has no effect in emitting pptx. Defaults to 'Normal'.
        """
        self.myrun = []
        self.current_styles = []
        self.current_list = []
        self.para = paragraph
        self.DBG = debug
        self.text_style = text_style
        super().__init__()

    def handle_starttag(self, tag, attrs):
        """Handles starting tags in rich text.

        Args:
            tag (str): The type of tag in the rich text.
            attrs (List[str]): Optional strings for extra formatting.
        """
        self.myrun.append(("start", tag))

    def handle_endtag(self, tag):
        """Handles ending tags in rich text.

        Args:
            tag (str): The type of tag in the rich text.
        """
        self.myrun.append(("end", tag))

    def handle_data(self, data):
        """Handles the data in between the tags of rich text.

        Args:
            data (str): The data inside of the tags.
        """
        self.myrun.append(("run", data))

    def reset_parser(self):
        """Empties the buffer inside of the current parser."""
        self.myrun = []

    def emit_docx(self):
        """Takes the current rich text inside of the parsers buffer and generates the appropriate actions for a docx document."""
        if self.DBG:
            print("-total run->", self.myrun)
        for (tag, action) in self.myrun:
            if self.DBG:
                print("|", tag, action)
            if tag == "start":
                if action == "p":
                    if self.DBG:
                        print("-emit->  p = self.doc.add_paragraph()")
                    p = self.para.insert_paragraph_before(style=self.text_style)
                elif action == "ul" or action == "ol":
                    self.current_list.append(self.list_types[action])
                    if self.DBG:
                        print("-adj-> appropriate list elements", self.current_list)
                elif action in self.style_eles:
                    self.current_styles.append(action)
                    if self.DBG:
                        print("-adj-> add style eles", self.current_styles)
                elif action == "li":
                    if self.DBG:
                        print(
                            "-emit-> p = self.doc.add_paragraph(style=",
                            self.current_list[-1],
                            ")",
                        )
                    p = self.para.insert_paragraph_before(style=self.current_list[-1])
                    if len(self.current_list) > 1:
                        list_level = len(self.current_list) - 1
                        if self.DBG:
                            print(
                                "-emit-> p.paragraph_format.left_indent = ... level=",
                                list_level,
                            )
                        p.paragraph_format.left_indent = Inches(
                            0.25 + (0.25 * list_level)
                        )
                elif action == "span":
                    self.current_styles.append(action)
            elif tag == "run":
                if self.DBG:
                    print("-emit-> r = p.add_run(", action, ")")
                r = p.add_run(action)
                if len(self.current_styles) != 0:
                    if self.DBG:
                        print("\tAdd styles", self.current_styles)
                    for cs in self.current_styles:
                        if cs == "b":
                            r.bold = True
                        if cs == "i":
                            r.italic = True
                        if cs == "u":
                            r.underline = True
            elif tag == "end":
                if self.DBG:
                    print("appropriate action for endtag:", action)
                if action == "span":
                    try:
                        self.current_styles.remove("b")
                    except:
                        self.current_styles.remove("span")
                elif action in self.style_eles:
                    self.current_styles.remove(action)
                    if self.DBG:
                        print("-adj-> current sytles in effect:", self.current_styles)
                elif action == "ul" or action == "ol":
                    self.current_list.pop()
                    if self.DBG:
                        print("-adj-> current list in effect:", self.current_list)

    def emit_pptx(self):
        """Takes the rich text inside of the parser and generates the appropriate pptx actions."""
        for (tag, action) in self.myrun:
            if tag == "start":
                if action == "p":
                    p = self.para.add_paragraph()
                elif action == "ul" or action == "ol":
                    self.current_list.append(self.list_types[action])
                elif action in self.style_eles:
                    self.current_styles.append(action)
                elif action == "li":
                    p = self.para.add_paragraph()
                    if len(self.current_list) > 0:
                        p.level = len(self.current_list)
                    if len(self.current_list) > 5:
                        p.level = 5
            elif tag == "run":
                r = p.add_run()
                r.text = action
                if len(self.current_styles) != 0:
                    for cs in self.current_styles:
                        if cs == "b":
                            r.font.bold = True
                        if cs == "i":
                            r.font.italic = True
                        if cs == "u":
                            r.font.underline = True
            elif tag == "end":
                if action == "span":
                    try:
                        self.current_styles.remove("b")
                    except:
                        self.current_styles.remove("span")
                elif action in self.style_eles:
                    self.current_styles.remove(action)
                elif action == "ul" or action == "ol":
                    self.current_list.pop()
