"""
This file contains various utilities to supplement python-docx for
advanced operations.

"""
# Risk & Vulnerability Assessment Reporting Engine

# Copyright 2022 The Risk & Vulnerability Reporting Engine Contributors, All Rights Reserved.
# (see Contributors.txt for a full list of Contributors)

# SPDX-License-Identifier: BSD-3-Clause

# Please see additional acknowledgments (including references to third party source code, object code, documentation and other files) in the license.txt file or contact permission@sei.cmu.edu for full terms.

# Created, in part, with funding and support from the United States Government. (see Acknowledgments file).

# DM22-1011

import sys
import os.path
import html
import datetime
import docx

from docx.oxml.ns import qn
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from lxml import etree
from PIL import Image

import report_gen.utilities.assessment_facts as af


def xpprint(e):
    """Pretty prints XML strings for debugging.

    Args:
        e (string): XML string to be printed out.
    """
    s = etree.tostring(e, pretty_print=True)
    print(s.decode('utf-8'))


def delete_paragraph(paragraph):
    """Takes a python docx paragraph object and deletes the its associated text from a docx document.

    Args:
        paragraph (python docx paragraph): The paragraph that will be deleted from the docx document.
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def move_table_after(table, paragraph):
    """Moves a docx table after another paragraph.

    Args:
        table (python docx table): The table to be moved after the paragraph.
        paragraph (python docx paragraph): The paragraph that is used to position the table.
    """
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def find_paragraph(doc, s):
    """Finds a paragraph in the specified docx document that contains the specified string of text.

    Args:
        doc (docx document): The docx document that will be searched for the specified string.
        s (string): The string in the docx document being searched.

    Returns:
        python-docx paragraph: The docx paragraph object that contains the text.
    """
    for para in doc.paragraphs:
        if s in para.text:
            return para
    return None


def xsafe(s):
    """Escapes HTML encoded text to prevent formatting issues.

    Args:
        s (string): HTML string that is to be escaped.

    Returns:
        string: The escaped html string.
    """
    return html.escape(s)


def dpi_safe(sfile, img):
    "Modify images to be DPI safe for docx to add_picture(), sfile is path"
    if img.info.get('dpi'):
        resave_flag = False
        dimensions = list(img.info.get('dpi'))
        for idx, dimension in enumerate(dimensions):
            if dimension == 0:
                dimensions[idx] = 72
                resave_flag = True
        if resave_flag:
            img.save(sfile, dpi=tuple(dimensions))


def set_column_width(column, width):
    """Takes a column of a python docx table and setting the width of each cell in that column.

    Args:
        column (python docx column): The column of cells that will have its width set.
        width (float): The size in inches that the column will be set to.
    """
    for cell in column.cells:
        cell.width = Inches(width)


def hex_to_rgb(hex):
    """Converts hex color codes and returns the converted RGB colors inside of a tuple.

    Args:
        hex (string): Hex color string

    Returns:
        tuple(int): Tuple contain the red, green, and blue values.
    """
    hex_strip = hex.lstrip('#')
    return tuple(int(hex_strip[i : i + 2], 16) for i in (0, 2, 4))


def insert_caption(doc, fig, text):
    """Inserts caption text for figures inside of the word document. This supports auto updating of numbering when updating the fields. The correct numbering
    will not appear at first until the user runs the auto update of all fields in the document using the word process engine.

    Args:
        doc (docx object): The word document
        fig (Boolean): True if caption belongs to a figure, False if caption belongs to a table
        text (String): The string that will be inserted under the image in the caption text.

    Returns:
        Docx Paragraph: The paragaph that the caption is inserted into.
    """
    if fig:
        para = doc.add_paragraph('Figure ', style='Caption')
    else:
        para = doc.add_paragraph('Table ', style='Caption')

    run = para.add_run()
    r = run._r

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')

    if fig:
        instrText.text = ' SEQ Figure \\* ARABIC'
    else:
        instrText.text = ' SEQ Table \\* ARABIC'

    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

    para.add_run(text=": " + text)
    
    return para

def resize_image(oldWidth, oldHeight, maxWidth, maxHeight):
    """
    Resizes an image while maintaining the same aspect ratio.
    All the arguments are casted to floats during resizing to assure compatability.
    The function checks to make sure that the new scaled dimensions to exceed either max boundry.
    Note: Pixel to Inches ratio is 72:1.

    Args:
        oldWidth (float): The original width of the image.
        oldHeight (float): The original height of the image.
        maxWidth (float): The maximum width the image can be.
        maxHeight (float): The maximum height the image can be.

    Returns:
        [(float,float)]: Returns new scaled width and height.
    """
    aspect_ratio = float(oldWidth) / float(oldHeight)

    if oldWidth < oldHeight:
        newWidth = float(maxHeight * aspect_ratio)
        if newWidth > maxWidth:
            return float(maxWidth), (float(maxWidth) / aspect_ratio)
        else:
            return (float(maxHeight) * aspect_ratio), float(maxHeight)
    else:
        newHeight = float(maxWidth) / aspect_ratio
        if newHeight > maxHeight:
            return (float(maxHeight) * aspect_ratio), float(maxHeight)
        else:
            return float(maxWidth), (float(maxWidth) / aspect_ratio)


def update_title(doc, db):
    """Updates and replaces content and tags that are located on the title page of the report.
    Since the title page puts its content inside of text boxes we have to search the XKL for the start of the textbox and the runs inside of it.
    It then searches and replaces all the tags in the text box.

    Args:
        doc (docx object): The report template
        db (List of dictionaries): The JSON data of the engagement.
    """

    emeta = af.get_db_info(db, "engagementmeta.fields", "keyNA", allow_empty=True)
    rep = af.get_db_info(db, "report.fields", "keyNA")

    body = doc._body._element
    text_box_p_elements = etree.ElementBase.xpath(
        body, './/wps:txbx//w:p', namespaces=body.nsmap
    )
    today = datetime.date.today()

    for ele in text_box_p_elements:
        # Since docx breaks up its text into multiple runs within a paragraph we have to search all of the runs to find the text we wanna replace.
        for run in ele.findall('.//' + qn('w:t')):
            # Empty textbox
            if run is None:
                continue

            #if '<STAKEHOLDER NAME>' in run.text:
            #    run.text = run.text.replace(
            #        '<STAKEHOLDER NAME>', emeta['customer_long_name']
            #    )
            if '{REPORT TITLE}' in run.text:
                run.text = run.text.replace(
                    '{REPORT TITLE}', rep['report_title']
                )
            if '{REPORT SUBTITLE}' in run.text:
                run.text = run.text.replace(
                    '{REPORT SUBTITLE}', rep['report_subtitle']
                )
            #if '[Serial Number]' in run.text:
            #    run.text = run.text.replace('[Serial Number]', emeta['asmt_id'])
            #if '<ASMT ID>' in run.text:
            #    run.text = run.text.replace(
            #        '<ASMT ID>', emeta['asmt_id']
            #    )
            if '{ DATE }' in run.text:
                run.text = run.text.replace(
                    '{ DATE }', today.strftime("%B %d, %Y"))


def update_charts(doc, media_path, p_tag, nist_file):
    """Updates the charts inside of the document.

    Args:
        doc (python docx document): The docx document that will be edited.
        media_path (string): The path to the media folder with screenshots.
        p_tag (string): The string inside of the document that will be updated.
        nist_file (string): The path to the chart file.
    """
    p_tag = find_paragraph(doc, p_tag)

    if p_tag is None:
        return  # tag must not be in the document

    # ---- create the ndf table
    p = p_tag._p  # starting point

    screen_p = doc.add_paragraph()
    #screen_p.style = 'Detailed Findings Normal'
    #screen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = screen_p.add_run()

    sfile = media_path + 'charts/' + nist_file
    if not os.path.exists(sfile):
        print("Can't find chart file", sfile)
        sys.exit(1)

    with Image.open(sfile) as img:
        swidth = img.size[0]

    if swidth < 455:
        r.add_picture(sfile)
    else:
        r.add_picture(sfile, width=Inches(6.5))

    p.addnext(screen_p._p)

    delete_paragraph(p_tag)


# Taken from https://programmersought.com/article/74085524416/
def set_cell_border(cell, **kwargs):
    """Allows the setting of specific borders inside of a table.
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )

    Args:
        cell (python docx table cell): The cell that will be edited.
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def force_update_fields(doc):
    """This function is designed to force the document to ask the user to update the fields in the document on open.
    NOTICE: This will cause a pop up to appear when they open the generated report. Clicking the OK button will not trigger anything malicious, but may catch them
    off guard if they are not expecitng it.

    Args:
        doc (docx object): The word document you are wanting to prompt for update.
    """
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    element_updatefields = etree.SubElement(
        doc.settings.element, f"{namespace}updateFields"
    )
    element_updatefields.set(f"{namespace}val", "true")


def add_link(paragraph, run, url, text):
    """This function will return a hyperlink to place in the document.

    Args:
        paragraph: the paragaph that the hyperlink will be inserted into
        run: the specific run that correlates with the hyperlink
        url: the url that the hyperlink will lead to
        text: the text that the hyperlink will be wrapped around
    """

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    paragraph.runs[run].font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    paragraph.runs[run].font.underline = True

    return hyperlink
