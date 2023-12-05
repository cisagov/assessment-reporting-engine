"""
The RVA report generator takes a document template and information
from the RVA pen testing django app database to generate an RVA report
in docx format
"""
# Risk & Vulnerability Assessment Reporting Engine

# Copyright 2022 The Risk & Vulnerability Reporting Engine Contributors, All Rights Reserved.
# (see Contributors.txt for a full list of Contributors)

# SPDX-License-Identifier: BSD-3-Clause

# Please see additional acknowledgments (including references to third party source code, object code, documentation and other files) in the license.txt file or contact permission@sei.cmu.edu for full terms.

# Created, in part, with funding and support from the United States Government. (see Acknowledgments file).

# DM22-1011

import sys
import os.path
import argparse
import datetime
import inflect

from datetime import timezone
from dateutil.relativedelta import relativedelta
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT

from PIL import Image

from .mam_xml import *  # mitre att&ck matrix
import report_gen.mam_xml as mam_xml
import report_gen.utilities.assessment_facts as af
import report_gen.utilities.xml_util as xu
import report_gen.utilities.rt_parser as rtp

try:
    import docx
except ImportError:
    print("Must have python docx library installed")
    sys.exit(1)


# ---- constants
tstamp = str(datetime.datetime.now().strftime("%Y%m%d_%H.%M.%S"))
today = datetime.date.today()
num = inflect.engine()


def insert_assessment_info(doc, db):
    """Function that fills out the tags on the report cover.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    asmt_id = af.get_db_info(db, 'engagementmeta.fields.asmt_id', '{ASMT ID}')
    stakeholder_name = af.get_db_info(db, 'engagementmeta.fields.customer_long_name', '{STAKEHOLDER NAME}')

    subtitle = "RV" + asmt_id + " - " + stakeholder_name

    p_tag = xu.find_paragraph(doc, "{REPORT SUBTITLE}")
    if p_tag is not None:
        p_tag.text = subtitle

    p_tag = xu.find_paragraph(doc, "{ DATE }")
    if p_tag is not None:
        p_tag.text = today.strftime("%B %-d, %Y")


def insert_fs_table(doc, db, fs_tag):
    """Generates and inserts the findings summary table.

    Args:
        doc (docx document): The docx template being edited.
        db (list): Assessment data
        fs_tag (string): The string in the docx document that will be replaced with the table.
    """

    # Find tag location in document
    p_tag = xu.find_paragraph(doc, fs_tag)

    # Stop if location does not exist
    if p_tag is None:
        return

    fs_table = doc.add_table(1, 4)
    fs_table.style = doc.styles['Findings Table']

    fs_table.cell(0, 0).text = ""
    fs_table.cell(0, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    fs_table.cell(0, 1).text = "Finding Name"
    fs_table.cell(0, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    fs_table.cell(0, 2).text = "Severity"
    fs_table.cell(0, 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    fs_table.cell(0, 3).text = "Mitigation Status"
    fs_table.cell(0, 3).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cnt, finding in enumerate(af.model_gen(db, "ptportal.uploadedfinding")):
        ele = finding['fields']
        row = fs_table.add_row()

        fid = str(cnt + 1)
        fname = ele['uploaded_finding_name']

        severity = ele['severity']

        if ele['mitigation']:
            mitigation = "Mitigated"
        else:
            mitigation = "Not Mitigated"

        fs_table.cell(cnt + 1, 0).text = fid
        fs_table.cell(cnt + 1, 0).paragraphs[0].runs[0].font.bold = True
        fs_table.cell(cnt + 1, 0).paragraphs[0].runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        fs_table.cell(cnt + 1, 1).text = fname
        fs_table.cell(cnt + 1, 1).paragraphs[0].runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        sev_cell = fs_table.cell(cnt + 1, 2).paragraphs[0]

        if severity == "Critical":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)
        elif severity == "High":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(252, 191, 143)
        elif severity == "Medium":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(255, 222, 89)
        elif severity == "Low":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
        else:
            sev_cell.add_run("• ").font.color.rgb = RGBColor(79, 175, 227)

        sev_cell.add_run(severity)
        fs_table.cell(cnt + 1, 2).paragraphs[0].runs[0].font.bold = True
        fs_table.cell(cnt + 1, 2).paragraphs[0].runs[1].font.bold = True
        fs_table.cell(cnt + 1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        mit_cell = fs_table.cell(cnt + 1, 3).paragraphs[0]

        if mitigation == "Mitigated":
            mit_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
        else:
            mit_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)

        mit_cell.add_run(mitigation)
        fs_table.cell(cnt + 1, 3).paragraphs[0].runs[0].font.bold = True
        fs_table.cell(cnt + 1, 3).paragraphs[0].runs[1].font.bold = True
        fs_table.cell(cnt + 1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    fs_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    fs_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    fs_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    fs_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    fs_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    fs_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
    fs_table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    fs_table.cell(0, 3).paragraphs[0].runs[0].font.bold = True

    xu.set_column_width(fs_table.columns[0], 0.38)
    xu.set_column_width(fs_table.columns[1], 3.38)
    xu.set_column_width(fs_table.columns[2], 1.38)
    xu.set_column_width(fs_table.columns[3], 1.37)

    xu.move_table_after(fs_table, p_tag)
    xu.delete_paragraph(p_tag)


def insert_kev_table(doc, db, kev_tag):
    """Generates and inserts the known exploited vulnerability table.

    Args:
        doc (docx document): The docx template being edited.
        db (list): Assessment data
        kev_tag (string): The string in the docx document that will be replaced with the table.
    """

    p_tag = xu.find_paragraph(doc, kev_tag)

    if p_tag is None:
        return

    kevs = []
    kev_ids = []

    for cnt, kev in enumerate(af.model_gen(db, "ptportal.kev")):
        ele = kev['fields']

        if ele['found']:
            found_kev = {"id": kev['pk'], "cve": ele['cve_id'], "vendor": ele['vendor_project'], "product": ele['product'], "vuln": ele['vulnerability_name'], "findings": []}
            kevs.append(found_kev)

    for cnt, finding in enumerate(af.model_gen(db, "ptportal.uploadedfinding")):
        ele = finding['fields']

        if ele['KEV']:        
            for item in kevs:
                if item['id'] in ele['KEV']:
                    item['findings'].append(ele['uploaded_finding_name'])

    kev_list = sorted(kevs, key=lambda d: d['cve']) 

    kev_table = doc.add_table(1, 5)
    kev_table.style = doc.styles['CISA Table']

    kev_table.cell(0, 0).text = "CVE"
    kev_table.cell(0, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 1).text = "Vendor"
    kev_table.cell(0, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 2).text = "Product"
    kev_table.cell(0, 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 3).text = "Vulnerability"
    kev_table.cell(0, 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 4).text = "Finding(s)"
    kev_table.cell(0, 4).alignment = WD_ALIGN_PARAGRAPH.CENTER

    kev_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 3).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 4).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 4).paragraphs[0].runs[0].font.bold = True

    for cnt, vuln in enumerate(kev_list):
        row = kev_table.add_row()

        url = "https://nvd.nist.gov/vuln/detail/" + vuln['cve']
        cve_cell = kev_table.cell(cnt + 1, 0)
        p = cve_cell.paragraphs[0]
        xu.add_link(p, 0, url, vuln['cve'])
        kev_table.cell(cnt + 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 1).text = vuln['vendor']
        kev_table.cell(cnt + 1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 2).text = vuln['product']
        kev_table.cell(cnt + 1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 3).text = vuln['vuln']
        kev_table.cell(cnt + 1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 4).text = ', '.join(vuln['findings'])
        kev_table.cell(cnt + 1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    xu.set_column_width(kev_table.columns[0], 0.6)
    xu.set_column_width(kev_table.columns[1], 0.72)
    xu.set_column_width(kev_table.columns[2], 0.81)
    xu.set_column_width(kev_table.columns[3], 2.19)
    xu.set_column_width(kev_table.columns[4], 2.19)

    xu.move_table_after(kev_table, p_tag)
    xu.delete_paragraph(p_tag)


def insert_cis_csc_table(doc, db, cis_csc_tag):
    """Generates and inserts the cis_csc controls table for the findings that are in the report.

    Args:
        doc (docx document): The docx template being edited.
        db (list): Assessment data
        cis_csc_tag (string): The string inside of the docx document that will be replaced with the table.
    """
    p_tag = xu.find_paragraph(doc, cis_csc_tag)

    if p_tag is None:
        return

    cis_csc_table = doc.add_table(1, 3)
    cis_csc_table.style = doc.styles['CISA Table']

    cis_csc_table.cell(0, 0).text = "ID"
    cis_csc_table.cell(0, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cis_csc_table.cell(0, 1).text = "Controls"
    cis_csc_table.cell(0, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    cis_csc_table.cell(0, 2).text = "Finding ID"
    cis_csc_table.cell(0, 2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cnt, rec in enumerate(af.model_gen(db, "ptportal.cis_csc")):
        ele = rec['fields']
        row = cis_csc_table.add_row()

        cis_csc_table.cell(cnt + 1, 0).text = str(ele['CIS_ID'])
        cis_csc_table.cell(cnt + 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cis_csc_table.cell(cnt + 1, 1).text = xu.xsafe(ele['name'])
        cis_csc_table.cell(cnt + 1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cis_csc_table.cell(cnt + 1, 2).text = ele['finding_ids']
        cis_csc_table.cell(cnt + 1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cis_csc_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cis_csc_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    cis_csc_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cis_csc_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    cis_csc_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cis_csc_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True

    xu.set_column_width(cis_csc_table.columns[0], 0.37)
    xu.set_column_width(cis_csc_table.columns[1], 3.5)
    xu.set_column_width(cis_csc_table.columns[2], 2.62)

    xu.move_table_after(cis_csc_table, p_tag)
    xu.delete_paragraph(p_tag)


def insert_df_table(doc, db, df_tag, media_path):
    """Generates the NCATS Detailed findings table for each finding that is listed/added to the report.

    Args:
        doc (docx document): The docx template being edited.
        db (list): Assessment data
        df_tag (string): The string inside of the docx document that will be replaced with the tables.
        media_path (string): Path to the location where the report screenshots are.
    """

    p_tag = xu.find_paragraph(doc, df_tag)

    if p_tag is None:
        return

    current_mark = p_tag

    # ---- find all screenshot information
    ss_info = af.build_screenshot_info(db)
    as_info = af.build_affected_systems_info(db)
    # make sure the affected systems are xml safe
    as_info = {k: xu.xsafe(v) for k, v in as_info.items()}

    for cnt, finding in enumerate(af.model_gen(db, "ptportal.uploadedfinding")):
        ele = finding['fields']

        fpk = finding['pk']
        fid = str(cnt + 1)
        fname = ele['uploaded_finding_name']

        severity = ele['severity']
        assessment_type = ele['assessment_type']

        if ele['mitigation']:
            mitigation = "Mitigated"
        else:
            mitigation = "Not Mitigated"

        finding_sshot = af.find_screenshots(ss_info, fpk)

        ptp_df_screen_text = "Relevant Screenshots"
        if len(finding_sshot) == 0:
            ptp_df_screenshot_note = "No relevant screenshots exist for this finding."
        else:
            ptp_df_screenshot_note = ele['screenshot_description']

        # ---- build the affected systems string
        affected_systems = af.find_affected_systems(as_info, ele['affected_systems'])

        df_table = doc.add_table(12, 5)
        df_table.style = doc.styles['Findings Table']

        df_table.cell(0, 0).text = "ID"
        df_table.cell(0, 1).text = "Finding"
        df_table.cell(0, 2).text = "Severity"
        df_table.cell(0, 3).text = "Type"
        df_table.cell(0, 4).text = "Mitigation Status"

        df_table.cell(1, 0).text = fid
        df_table.cell(1, 1).text = xu.xsafe(fname)
        #df_table.cell(1, 2).text = "• " + xu.xsafe(severity)
        sev_cell = df_table.cell(1, 2).paragraphs[0]
        
        if severity == "Critical":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)
        elif severity == "High":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(252, 191, 143)
        elif severity == "Medium":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(255, 222, 89)
        elif severity == "Low":
            sev_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
        else:
            sev_cell.add_run("• ").font.color.rgb = RGBColor(79, 175, 227)

        sev_cell.add_run(xu.xsafe(severity))
        df_table.cell(1, 2).paragraphs[0].runs[0].font.bold = True
        df_table.cell(1, 2).paragraphs[0].runs[1].font.bold = True
        df_table.cell(1, 3).text = xu.xsafe(assessment_type)

        mit_cell = df_table.cell(1, 4).paragraphs[0]

        if mitigation == "Mitigated":
            mit_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
        else:
            mit_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)

        mit_cell.add_run(xu.xsafe(mitigation))
        df_table.cell(1, 4).paragraphs[0].runs[0].font.bold = True
        df_table.cell(1, 4).paragraphs[0].runs[1].font.bold = True

        for i in range(0, 5):
            df_table.cell(0, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            df_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            xu.set_cell_border(df_table.cell(0, i), top={"sz": 8, "val": "single", "color": "#000000"})
            xu.set_cell_border(df_table.cell(0, i), bottom={"sz": 8, "val": "single", "color": "#000000"})
            if i < 4:
                xu.set_cell_border(df_table.cell(0, i), end={"sz": 8, "val": "single", "color": "#000000"})

        for i in range(0, 5):
            df_table.cell(1, i).paragraphs[0].runs[0].font.bold = True
            xu.set_cell_border(df_table.cell(1, i), top={"sz": 8, "val": "single", "color": "#000000"})
            xu.set_cell_border(df_table.cell(1, i), bottom={"sz": 8, "val": "single", "color": "#000000"})
            if i < 4:
                xu.set_cell_border(df_table.cell(1, i), end={"sz": 8, "val": "single", "color": "#000000"})

        df_table.cell(2, 0).merge(df_table.cell(2, 4))
        df_table.cell(2, 0).text = "Affected Systems"
        df_table.cell(2, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(3, 82, 136)
        df_table.cell(2, 0).paragraphs[0].runs[0].font.bold = True
        df_table.cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        xu.set_cell_border(df_table.cell(2, 0), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
        df_table.cell(3, 0).merge(df_table.cell(3, 4))
        df_table.cell(3, 0).text = affected_systems
        df_table.cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        df_table.cell(4, 0).merge(df_table.cell(4, 4))
        df_table.cell(4, 0).text = "Description"
        df_table.cell(4, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(3, 82, 136)
        df_table.cell(4, 0).paragraphs[0].runs[0].font.bold = True
        df_table.cell(4, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        xu.set_cell_border(df_table.cell(4, 0), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
        df_table.cell(5, 0).merge(df_table.cell(5, 4))
        df_table.cell(5, 0).text = xu.xsafe(ele['description'])
        df_table.cell(5, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        df_table.cell(6, 0).merge(df_table.cell(6, 4))
        df_table.cell(6, 0).text = "Recommended Mitigation"
        df_table.cell(6, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(3, 82, 136)
        df_table.cell(6, 0).paragraphs[0].runs[0].font.bold = True
        df_table.cell(6, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        xu.set_cell_border(df_table.cell(6, 0), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
        df_table.cell(7, 0).merge(df_table.cell(7, 4))
        df_table.cell(7, 0).text = xu.xsafe(ele['remediation'])
        df_table.cell(7, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        df_table.cell(8, 0).merge(df_table.cell(8, 4))
        df_table.cell(8, 0).text = xu.xsafe(ptp_df_screen_text)
        df_table.cell(8, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(3, 82, 136)
        df_table.cell(8, 0).paragraphs[0].runs[0].font.bold = True
        df_table.cell(8, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        xu.set_cell_border(df_table.cell(8, 0), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
        df_table.cell(9, 0).merge(df_table.cell(9, 4))
        df_table.cell(9, 0).paragraphs[0].text = ptp_df_screenshot_note
        df_table.cell(9, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        df_table.cell(10, 0).merge(df_table.cell(10, 4))
        df_table.cell(10, 0).text = "Security References"
        df_table.cell(10, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(3, 82, 136)
        df_table.cell(10, 0).paragraphs[0].runs[0].font.bold = True
        df_table.cell(10, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        xu.set_cell_border(df_table.cell(10, 0), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
        df_table.cell(11, 0).merge(df_table.cell(11, 4))

        if len(ele['KEV']) > 0:
            kev = df_table.cell(11, 0).paragraphs[0]
            kev.add_run("This finding is a ")
            kev.add_run("Known Exploited Vulnerability.\n").bold = True

        nist_800 = df_table.cell(11, 0).paragraphs[0]
        nist_800.add_run("NIST 800-53: ").bold = True
        nist_800.add_run(str(ele['NIST_800_53']) + '\n')
        nist_csf = df_table.cell(11, 0).paragraphs[0]
        nist_csf.add_run("NIST CSF: ").bold = True
        nist_csf.add_run(str(ele['NIST_CSF']))
        df_table.cell(11, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        xu.set_column_width(df_table.columns[0], 0.42)
        xu.set_column_width(df_table.columns[1], 2.22)
        xu.set_column_width(df_table.columns[2], 1.23)
        xu.set_column_width(df_table.columns[3], 1.25)
        xu.set_column_width(df_table.columns[4], 1.32)

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        current_mark._p.addnext(page_break._p)
        xu.move_table_after(df_table, current_mark)

        screenshot_cell = df_table.cell(9, 0)
        for sshot in finding_sshot:
            ssf = sshot['fields']
            sfile = media_path + ssf['file']
            if not os.path.exists(sfile):
                print("Can't find screenshot file", sfile)
                continue

            screenshot = screenshot_cell.add_paragraph()
            screenshot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            screenshot_run = screenshot.add_run()

            scap = xu.xsafe(ssf['caption'])

            with Image.open(sfile) as img:
                swidth = img.size[0]
                xu.dpi_safe(sfile, img)

            if swidth < 455:
                screenshot_run.add_picture(sfile)
            else:
                screenshot_run.add_picture(sfile, width=Inches(5.5))

            caption = xu.insert_caption(doc, True, scap)
            screenshot._p.addnext(caption._p)

        # now delete the temporary screenshot xml tag
        current_mark = page_break

    # ---- remove the paragraph containing the tag
    xu.delete_paragraph(p_tag)


def insert_rs_table(doc, db, rs_tag, media_path):
    """Generates and inserts the table with the risk scores and findings chart.

    Args:
        doc (docx document): The docx template being edited.
        db (list): Assessment data
        rs_tag (string): The string in the docx document that will be replaced with the table.
        media_path (string): Path to the location where the report images are.
    """

    p_tag = xu.find_paragraph(doc, rs_tag)

    if p_tag is None:
        return

    total_risk_score = 0
    mitigated_risk_score = 0

    for finding in af.model_gen(db, 'ptportal.uploadedfinding'):
        total_risk_score += finding.get("fields").get("risk_score")
        if not finding.get("fields").get("mitigation"):
            mitigated_risk_score += finding.get("fields").get("risk_score")

    rs_table = doc.add_table(1, 2)
    score_cell = rs_table.cell(0, 0)
    chart_cell = rs_table.cell(0, 1)

    total_label = score_cell.add_paragraph()
    total_label.text = "Total Risk Score"
    total_label.style = "NormalBold"
    total_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_score = score_cell.add_paragraph()
    total_score.text = str(total_risk_score)
    total_score.style = "RiskScore"

    mitigated_label = score_cell.add_paragraph()
    mitigated_label.text = "Mitigated Risk Score"
    mitigated_label.style = "NormalBold"
    mitigated_label.alignment = WD_ALIGN_PARAGRAPH.CENTER

    mitigated_score = score_cell.add_paragraph()
    mitigated_score.text = str(mitigated_risk_score)
    mitigated_score.style = "RiskScore"

    blank = score_cell.add_paragraph()
    blank.text = ""
    blank.alignment = WD_ALIGN_PARAGRAPH.CENTER

    file = media_path + 'charts/riskchart.png'
    if not os.path.exists(file):
        print("Can't find chart file", file)
    else:
        chart = chart_cell.paragraphs[0].add_run()
        chart_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        with Image.open(file) as img:
            swidth = img.size[0]
        if swidth < 455:
            chart.add_picture(file)
        else:
            chart.add_picture(file, width=Inches(4.5))

    xu.set_cell_border(rs_table.cell(0, 0), start={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 1), start={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 0), top={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 1), top={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 0), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 1), bottom={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 0), end={"sz": 8, "val": "single", "color": "#BFBFBF"})
    xu.set_cell_border(rs_table.cell(0, 1), end={"sz": 8, "val": "single", "color": "#BFBFBF"})

    xu.set_column_width(rs_table.columns[0], 1.62)
    xu.set_column_width(rs_table.columns[1], 4.87)

    xu.move_table_after(rs_table, p_tag)
    
    xu.delete_paragraph(p_tag)


def insert_pass_analysis(doc, db):
    """Inserts the password analysis file into the appendix.

    Args:
        doc (Docx object): The docx template.
        db (List of dictionaries): Json of the engagement data.
    """
    p_tag = xu.find_paragraph(doc, "{PASSWORD ANALYSIS}")

    if p_tag is None:
        return

    p_tag.style = doc.styles['Pass Analysis']
    p_tag.text = af.get_db_info(db, 'report.fields.password_analysis', 'If password analysis data is available, paste it here. Otherwise, consider removing this section and renaming the subsequent appendices accordingly.')


def insert_acronyms(doc, db):
    """Inserts list of acronyms into the appendix.

    Args:
        doc (Docx object): The docx template.
        db (List of dictionaries): Json of the engagement data.
    """

    p_tag = xu.find_paragraph(doc, "{Table: Acronyms}")

    if p_tag is None:
        return

    acronym_table = doc.add_table(1, 2)
    acronym_table.style = doc.styles['Acronyms']

    acronyms = []

    for cnt, acr in enumerate(af.model_gen(db, "ptportal.acronym")):
        ele = acr['fields']
        acronym = {}
        acronym["abbr"] = ele['acronym']
        acronym["def"] = ele['definition']
        acronyms.append(acronym)

    org_acronym = {}
    org_acronym["abbr"] = af.get_db_info(db, 'engagementmeta.fields.customer_initials', '{STAKEHOLDER ACRONYM}')
    org_acronym["def"] = af.get_db_info(db, 'engagementmeta.fields.customer_long_name', '{STAKEHOLDER NAME}')
    acronyms.append(org_acronym)
    sorted_acronyms = sorted(acronyms, key=lambda x: x['abbr'])

    for cnt, acr in enumerate(sorted_acronyms):
        row = acronym_table.add_row()

        acronym_table.cell(cnt, 0).text = acr["abbr"]
        acronym_table.cell(cnt, 1).text = acr["def"]
        
    xu.set_column_width(acronym_table.columns[0], 1.38)
    xu.set_column_width(acronym_table.columns[1], 5.23)

    xu.move_table_after(acronym_table, p_tag)
    
    xu.delete_paragraph(p_tag)


def insert_attack_paths(doc, db, media_path):
    """Inserts all the attack paths that were uploaded and inserts them into the template. Function centers images and adds a page break after every single one.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
        media_path (String): Path to the media files
    """

    p_tag = xu.find_paragraph(doc, "{ATTACK PATHS}")

    techniques = []

    for technique in af.model_gen(db, "ptportal.attack"):
        ele = technique['fields']
        t_data = {"id": technique['pk'], "t_id": ele['t_id'], "name": ele['name'], "url": ele['url']}
        techniques.append(t_data)

    for cnt, path in reversed(list(enumerate(af.model_gen(db, "ptportal.narrative")))):

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

        ele = path['fields']

        name = ele['name'] + " " + str(ele['order'])
        used_techniques = []

        if ele['attack']:
            for item in techniques:
                if item['id'] in ele['attack']:
                    used_techniques.append(item)

        for t in reversed(used_techniques):
            b = doc.add_paragraph()
            if t['url']:
                url = t['url']
            else:
                url = "https://attack.mitre.org/techniques/"
            xu.add_link(b, 0, url, t['name'])
            b.style = "List Bullet 2"
            p_tag._p.addnext(b._p)

        mitre = doc.add_paragraph()
        mitre_start = mitre.add_run("The following ")
        xu.add_link(mitre, 1, "https://attack.mitre.org/", "MITRE ATT&CK")
        mitre_end = mitre.add_run(" Techniques were leveraged in this attack path:")
        mitre.runs[1].font.bold = True
        mitre.style = "Normal"
        p_tag._p.addnext(mitre._p)

        if ele['file']:
            dcap = xu.xsafe(ele['caption'])
            caption = xu.insert_caption(doc, True, dcap)
            p_tag._p.addnext(caption._p)

            file = media_path + ele['file']
            if not os.path.exists(file):
                print("Can't find diagram file", file)
                continue
            diagram = doc.add_paragraph()
            diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
            diagram_run = diagram.add_run()

            with Image.open(file) as img:
                swidth = img.size[0]
                xu.dpi_safe(file, img)

            if swidth < 455:
                diagram_run.add_picture(file)
            else:
                diagram_run.add_picture(file, width=Inches(6.5))

            p_tag._p.addnext(diagram._p)
        else:
            diagram = doc.add_paragraph("<NO ATTACK PATH DIAGRAM>")
            diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print("No diagram.")

        header = doc.add_paragraph(name, style="Heading 2")
        p_tag._p.addnext(header._p)

    xu.delete_paragraph(p_tag)


def insert_narrative(doc, db, media_path):
    """Fills out the narrative section of the report.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
        media_path (String): Path to the media files.
    """

    p_tag = xu.find_paragraph(doc, "{NARRATIVE SECTION}")

    if p_tag is None:
        return

    narratives = af.model_gen(db, 'ptportal.narrative')

    external_narratives = []
    internal_narratives = []
    phishing_narratives = []
    tools = []

    for cnt, tool in enumerate(af.model_gen(db, 'ptportal.tools')):
        ele = tool['fields']
        t = {"id": tool['pk'], "name": ele['name'], "url": ele['url']}
        tools.append(t)

    for narrative in narratives:
        ele = narrative['fields']

        if ele['assessment_type'] == 1:
            external_narratives.append(narrative)
        elif ele['assessment_type'] == 2:
            internal_narratives.append(narrative)
        else:
            phishing_narratives.append(narrative)

    if len(phishing_narratives) > 0:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)
        
        for p_cnt, pn in reversed(list(enumerate(phishing_narratives))):
            narrative_steps = af.model_gen(db, 'ptportal.narrativestep')

            for pc, p_step in reversed(list(enumerate(narrative_steps))):
                ele = p_step['fields']
                if pn['pk'] == ele['narrative']:
                    if ele['file']:
                        p_scap = xu.xsafe(ele['caption'])
                        p_caption = xu.insert_caption(doc, True, p_scap)
                        p_tag._p.addnext(p_caption._p)

                        file = media_path + ele['file']
                        if not os.path.exists(file):
                            print("Can't find diagram file", file)
                            continue
                        p_ss = doc.add_paragraph()
                        p_ss.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_ss_run = p_ss.add_run()

                        with Image.open(file) as img:
                            swidth = img.size[0]
                            xu.dpi_safe(file, img)

                        if swidth < 455:
                            p_ss_run.add_picture(file)
                        else:
                            p_ss_run.add_picture(file, width=Inches(5.5))

                        p_tag._p.addnext(p_ss._p)

                    p_description = doc.add_paragraph(ele['description'], style="Normal")
                    p_tag._p.addnext(p_description._p)

            if pn['fields']['tools']:
                p_tools = []

                tool_space = doc.add_paragraph("", style="Normal")
                p_tag._p.addnext(tool_space._p)
                
                for item in tools:
                    if item['id'] in pn['fields']['tools']:
                        p_tools.append(item)
                for tool in reversed(p_tools):
                    b = doc.add_paragraph()
                    tool_name = b.add_run(tool['name'] + " ")
                    if tool['url']:
                        ob = b.add_run("[")
                        tool_url = b.add_run()
                        xu.add_link(b, 2, tool['url'], tool['url'])
                        cb = b.add_run("]")
                    b.style = "List Bullet 2"
                    p_tag._p.addnext(b._p)

                tool_header = doc.add_paragraph("Relevant Tools:", style="Normal")
                p_tag._p.addnext(tool_header._p)

            p_ap_header = doc.add_paragraph("Attack Path " + str(p_cnt + 1), style="Heading 3")
            p_tag._p.addnext(p_ap_header._p)

        phi_header = doc.add_paragraph("Phishing", style="Heading 2")
        p_tag._p.addnext(phi_header._p)


    if len(internal_narratives) > 0:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)
        
        for i_cnt, inn in reversed(list(enumerate(internal_narratives))):

            narrative_steps = af.model_gen(db, 'ptportal.narrativestep')

            for ic, i_step in reversed(list(enumerate(narrative_steps))):
                ele = i_step['fields']
                if inn['pk'] == ele['narrative']:
                    if ele['file']:
                        i_scap = xu.xsafe(ele['caption'])
                        i_caption = xu.insert_caption(doc, True, i_scap)
                        p_tag._p.addnext(i_caption._p)

                        file = media_path + ele['file']
                        if not os.path.exists(file):
                            print("Can't find diagram file", file)
                            continue
                        i_ss = doc.add_paragraph()
                        i_ss.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        i_ss_run = i_ss.add_run()

                        with Image.open(file) as img:
                            swidth = img.size[0]
                            xu.dpi_safe(file, img)

                        if swidth < 455:
                            i_ss_run.add_picture(file)
                        else:
                            i_ss_run.add_picture(file, width=Inches(5.5))

                        p_tag._p.addnext(i_ss._p)

                    i_description = doc.add_paragraph(ele['description'], style="Normal")
                    p_tag._p.addnext(i_description._p)

            if inn['fields']['tools']:
                i_tools = []
                
                for item in tools:
                    if item['id'] in inn['fields']['tools']:
                        i_tools.append(item)
                for tool in reversed(i_tools):
                    b = doc.add_paragraph()
                    tool_name = b.add_run(tool['name'] + " ")
                    if tool['url']:
                        ob = b.add_run("[")
                        tool_url = b.add_run()
                        xu.add_link(b, 2, tool['url'], tool['url'])
                        cb = b.add_run("]")
                    b.style = "List Bullet 2"
                    p_tag._p.addnext(b._p)

                tool_header = doc.add_paragraph("Relevant Tools:", style="Normal")
                p_tag._p.addnext(tool_header._p)

            i_ap_header = doc.add_paragraph("Attack Path " + str(i_cnt + 1), style="Heading 3")
            p_tag._p.addnext(i_ap_header._p)

        int_header = doc.add_paragraph("Internal", style="Heading 2")
        p_tag._p.addnext(int_header._p)


    if len(external_narratives) > 0:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)
        
        for e_cnt, en in reversed(list(enumerate(external_narratives))):

            narrative_steps = af.model_gen(db, 'ptportal.narrativestep')

            for ec, e_step in reversed(list(enumerate(narrative_steps))):
                ele = e_step['fields']
                if en['pk'] == ele['narrative']:
                    if ele['file']:
                        e_scap = xu.xsafe(ele['caption'])
                        e_caption = xu.insert_caption(doc, True, e_scap)
                        p_tag._p.addnext(e_caption._p)

                        file = media_path + ele['file']
                        if not os.path.exists(file):
                            print("Can't find diagram file", file)
                            continue
                        e_ss = doc.add_paragraph()
                        e_ss.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        e_ss_run = e_ss.add_run()

                        with Image.open(file) as img:
                            swidth = img.size[0]
                            xu.dpi_safe(file, img)

                        if swidth < 455:
                            e_ss_run.add_picture(file)
                        else:
                            e_ss_run.add_picture(file, width=Inches(5.5))

                        p_tag._p.addnext(e_ss._p)

                    e_description = doc.add_paragraph(ele['description'], style="Normal")
                    p_tag._p.addnext(e_description._p)

            if en['fields']['tools']:
                e_tools = []
                
                for item in tools:
                    if item['id'] in en['fields']['tools']:
                        e_tools.append(item)
                for tool in reversed(e_tools):
                    b = doc.add_paragraph()
                    tool_name = b.add_run(tool['name'] + " ")
                    if tool['url']:
                        ob = b.add_run("[")
                        tool_url = b.add_run()
                        xu.add_link(b, 2, tool['url'], tool['url'])
                        cb = b.add_run("]")
                    b.style = "List Bullet 2"
                    p_tag._p.addnext(b._p)

                tool_header = doc.add_paragraph("Relevant Tools:", style="Normal")
                p_tag._p.addnext(tool_header._p)

            e_ap_header = doc.add_paragraph("Attack Path " + str(e_cnt + 1), style="Heading 3")
            p_tag._p.addnext(e_ap_header._p)

        ext_header = doc.add_paragraph("External", style="Heading 2")
        p_tag._p.addnext(ext_header._p)

    xu.delete_paragraph(p_tag)


def insert_report_summary(doc, db, media_path):
    """Function that fills out all the tags inside of the summary section of the report.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
        media_path (String): Path to the media files
    """

    nc_tag = xu.find_paragraph(doc, "{NIST CONTROLS}")

    if nc_tag is not None:
        file = media_path + 'charts/nistspchart.png'
        if not os.path.exists(file):
            print("Can't find chart file", file)
        else:
            para = doc.add_paragraph()
            r = para.add_run()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            with Image.open(file) as img:
                swidth = img.size[0]
            if swidth < 455:
                r.add_picture(file)
            else:
                r.add_picture(file, width=Inches(6.5))

            nc_tag._p.addnext(para._p)
            xu.delete_paragraph(nc_tag)

    nf_tag = xu.find_paragraph(doc, "{NIST CSF}")

    if nf_tag is not None:
        file = media_path + 'charts/nistframeworkchart.png'
        if not os.path.exists(file):
            print("Can't find chart file", file)
        else:
            para = doc.add_paragraph()
            r = para.add_run()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            with Image.open(file) as img:
                swidth = img.size[0]
            if swidth < 455:
                r.add_picture(file)
            else:
                r.add_picture(file, width=Inches(6.5))

            nf_tag._p.addnext(para._p)
            xu.delete_paragraph(nf_tag)

    # Creates the bulleted sections in the summary.
    bullet_lists = [
        ('{SIGNIFICANT FINDINGS}', 'report.fields.significant_findings'),
        ('{ RECOMMENDATIONS }', 'report.fields.recommendations'),
        ('{OBSERVED STRENGTHS}', 'report.fields.observed_strengths'),
    ]
    for bl in bullet_lists:
        p_tag = xu.find_paragraph(doc, bl[0])
        if p_tag is not None:
            text = af.get_db_info(db, bl[1], "").replace("\r", "\n")
            for bullet in text.split("\n")[::-1]:
                para = doc.add_paragraph()
                para.text = bullet
                para.style = "List Bullet 2"
                p_tag._p.addnext(para._p)
            xu.delete_paragraph(p_tag)

    # Counts the number of findings of the severities.
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Informational": 0}
    words = {"Critical": "", "High": "", "Medium": "", "Low": "", "Informational": ""}
    plural = {"Critical": "findings", "High": "findings", "Medium": "findings", "Low": "findings", "Informational": "findings"}
    total_risk_score = 0
    mitigated_risk_score = 0
    for finding in af.model_gen(db, 'ptportal.uploadedfinding'):
        counts[finding.get("fields").get("severity")] = (
            counts.get(finding.get("fields").get("severity"), 0) + 1
        )
        total_risk_score += finding.get("fields").get("risk_score")
        if not finding.get("fields").get("mitigation"):
            mitigated_risk_score += finding.get("fields").get("risk_score")

    words['Critical'] = num.number_to_words(counts['Critical'])
    words['High'] = num.number_to_words(counts['High'])
    words['Medium'] = num.number_to_words(counts['Medium'])
    words['Low'] = num.number_to_words(counts['Low'])
    words['Informational'] = num.number_to_words(counts['Informational'])

    for sev in counts:
        if counts[sev] == 1:
            plural[sev] = "finding"

    p_tag = xu.find_paragraph(doc, "{FINDINGS SUMMARY}")
    if p_tag is not None:
        p_tag.text = f"CISA identified {words['Critical']} ({counts['Critical']}) critical {plural['Critical']}, {words['High']} ({counts['High']}) high {plural['High']}, {words['Medium']} ({counts['Medium']}) medium {plural['Medium']}, {words['Low']} ({counts['Low']}) low {plural['Low']}, and {words['Informational']} ({counts['Informational']}) informational {plural['Informational']}. The Total Risk Score for this assessment is {total_risk_score} and the Mitigated Risk Score after deducting mitigated findings is {mitigated_risk_score}."

    p_tag = xu.find_paragraph(doc, "{Total Score}")
    if p_tag is not None:
        p_tag.text = str(total_risk_score)

    p_tag = xu.find_paragraph(doc, "{Mitigated Score}")
    if p_tag is not None:
        p_tag.text = str(mitigated_risk_score)


def insert_scope(doc, db):
    """Function that fills out all the tags inside of the scope section of the report.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{PHISHING USERS}")
    if p_tag is not None:
        phishing_users = af.get_db_info(db, 'report.fields.users_targeted', '{USERS TARGETED}')
        p_tag.text = f"{phishing_users} users were included in the phishing campaign(s) targeting the following mail domain(s):"

    p_tag = xu.find_paragraph(doc, "{PHISHING DOMAINS}")
    if p_tag is not None:
        p_tag.text = af.get_db_info(db, 'engagementmeta.fields.phishing_domains', '{PHISHING DOMAINS}')

    p_tag = xu.find_paragraph(doc, "{EXTERNAL HOSTS}")
    if p_tag is not None:
        ext_disc = af.get_db_info(db, 'report.fields.external_discovered', '{EXTERNAL DISCOVERED}')
        ext_scan = af.get_db_info(db, 'report.fields.external_scanned', '{EXTERNAL SCANNED}')
        p_tag.text = f"{ext_disc} of {ext_scan} scanned hosts were discovered to be live during testing of the following in scope targets:"

    p_tag = xu.find_paragraph(doc, "{EXTERNAL SCOPE}")
    if p_tag is not None:
        p_tag.text = af.get_db_info(db, 'engagementmeta.fields.ext_scope', '{EXTERNAL SCOPE}')

    p_tag = xu.find_paragraph(doc, "{INTERNAL HOSTS}")
    if p_tag is not None:
        int_disc = af.get_db_info(db, 'report.fields.internal_discovered', '{INTERNAL DISCOVERED}')
        int_scan = af.get_db_info(db, 'report.fields.internal_scanned', '{INTERNAL SCANNED}')
        p_tag.text = f"{int_disc} of {int_scan} scanned hosts were discovered to be live during testing of the following in scope targets:"

    p_tag = xu.find_paragraph(doc, "{INTERNAL SCOPE}")
    if p_tag is not None:
        p_tag.text =af.get_db_info(db, 'engagementmeta.fields.int_scope', '{INTERNAL SCOPE}')

    p_tag = xu.find_paragraph(doc, "{OSINF DOMAINS}")
    if p_tag is not None:
        p_tag.text =af.get_db_info(db, 'engagementmeta.fields.osinf_scope', '{OSINF DOMAINS}')

    p_tag = xu.find_paragraph(doc, "{WEB APP SCOPE}")
    if p_tag is not None:
        p_tag.text =af.get_db_info(db, 'engagementmeta.fields.web_app_scope', '{WEB APP SCOPE}')


def insert_ransomware(doc, db):
    """Function that fills out the ransomware susceptibility section of the report, if ransomware results exist.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{RANSOMWARE RESULTS}")

    if p_tag is None:
        return

    if af.get_db_info(db, 'ransomwarescenarios', 'NA') is not None or af.get_db_info(db, 'ransomware', 'NA') is not None:

        rg_url = "https://www.cisa.gov/stopransomware/ransomware-guide"
        rr_url = "https://www.cisa.gov/stopransomware/report-ransomware-0"

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

        respond = doc.add_paragraph()
        r_run1 = respond.add_run("In the event that a ransomware attack has taken place, the incident should be reported to federal law enforcement (see ")
        xu.add_link(respond, 1, rr_url, "CISA's ransomware report page")
        r_run3 = respond.add_run(") and it is recommended to follow the steps outlined in the ")
        r_run4 = respond.add_run("Ransomware Response Checklist")
        r_run5 = respond.add_run(" section of the ")
        xu.add_link(respond, 5, rg_url, "CISA Ransomware Guide")
        r_end = respond.add_run(".")
        respond.runs[3].font.bold = True
        respond.style = "Normal"
        p_tag._p.addnext(respond._p)

        prevent = doc.add_paragraph()
        p_run1 = prevent.add_run("In addition to the recommended mitigations provided throughout the ")
        p_run2 = prevent.add_run("Findings")
        p_run3 = prevent.add_run(" section of this report, the CISA team recommends taking the proactive actions listed in the ")
        p_run4 = prevent.add_run("Ransomware Prevention Best Practices")
        p_run5 = prevent.add_run(" section of the ")
        xu.add_link(prevent, 5, rg_url, "CISA Ransomware Guide")
        p_end = prevent.add_run(".")
        prevent.runs[1].font.bold = True
        prevent.runs[3].font.bold = True
        prevent.style = "Normal"
        p_tag._p.addnext(prevent._p)

        if af.get_db_info(db, 'ransomware', 'NA') is not None:
            ransomware_results = []

            for item in af.model_gen(db, 'ptportal.ransomware'):
                ele = item['fields']
                if not ele['disabled']:
                    if "detected by security software" in ele['description']:
                        if ele['trigger'] == "Y":
                            if ele['time_start'] and ele['time_end']:
                                start = datetime.datetime.fromisoformat(ele['time_start'][:-1]).astimezone(timezone.utc)
                                end = datetime.datetime.fromisoformat(ele['time_end'][:-1]).astimezone(timezone.utc)
                                difference = relativedelta(end, start)

                                if difference.days == 1:
                                    days = "1 day "
                                else:
                                    days = str(difference.days) + " days "
                                if difference.hours == 1:
                                    hours = "1 hour "
                                else:
                                    hours = str(difference.hours) + " hours "
                                if difference.minutes == 1:
                                    minutes = "and 1 minute."
                                else:
                                    minutes = "and " + str(difference.minutes) + " minutes."
                            else:
                                days = "0 days "
                                hours = "0 hours "
                                minutes = "and 0 minutes."
                            ransomware_results.append("Ransomware activity was detected by security software within " + days + hours + minutes)
                        else:
                            ransomware_results.append("Ransomware activity was not detected by security software.")

                    if "prevented by security software" in ele['description']:
                        if ele['trigger'] == "Y":
                            if ele['time_start'] and ele['time_end']:
                                start = datetime.datetime.fromisoformat(ele['time_start'][:-1]).astimezone(timezone.utc)
                                end = datetime.datetime.fromisoformat(ele['time_end'][:-1]).astimezone(timezone.utc)
                                difference = relativedelta(end, start)

                                if difference.days == 1:
                                    days = "1 day "
                                else:
                                    days = str(difference.days) + " days "
                                if difference.hours == 1:
                                    hours = "1 hour "
                                else:
                                    hours = str(difference.hours) + " hours "
                                if difference.minutes == 1:
                                    minutes = "and 1 minute."
                                else:
                                    minutes = "and " + str(difference.minutes) + " minutes."
                            else:
                                days = "0 days "
                                hours = "0 hours "
                                minutes = "and 0 minutes."
                            ransomware_results.append("Ransomware activity was prevented by security software within " + days + hours + minutes)
                        else:
                            ransomware_results.append("Ransomware activity was not prevented by security software.")

                    if "detected by security and/or IT personnel" in ele['description']:
                        if ele['trigger'] == "Y":
                            if ele['time_start'] and ele['time_end']:
                                start = datetime.datetime.fromisoformat(ele['time_start'][:-1]).astimezone(timezone.utc)
                                end = datetime.datetime.fromisoformat(ele['time_end'][:-1]).astimezone(timezone.utc)
                                difference = relativedelta(end, start)

                                if difference.days == 1:
                                    days = "1 day "
                                else:
                                    days = str(difference.days) + " days "
                                if difference.hours == 1:
                                    hours = "1 hour "
                                else:
                                    hours = str(difference.hours) + " hours "
                                if difference.minutes == 1:
                                    minutes = "and 1 minute."
                                else:
                                    minutes = "and " + str(difference.minutes) + " minutes."
                            else:
                                days = "0 days "
                                hours = "0 hours "
                                minutes = "and 0 minutes."
                            ransomware_results.append("Ransomware activity was detected by security and/or IT personnel within " + days + hours + minutes)
                        else:
                            ransomware_results.append("Ransomware activity was not detected by security and/or IT personnel.")

                    if "reported by end users" in ele['description']:
                        if ele['trigger'] == "Y":
                            if ele['time_start'] and ele['time_end']:
                                start = datetime.datetime.fromisoformat(ele['time_start'][:-1]).astimezone(timezone.utc)
                                end = datetime.datetime.fromisoformat(ele['time_end'][:-1]).astimezone(timezone.utc)
                                difference = relativedelta(end, start)

                                if difference.days == 1:
                                    days = "1 day "
                                else:
                                    days = str(difference.days) + " days "
                                if difference.hours == 1:
                                    hours = "1 hour "
                                else:
                                    hours = str(difference.hours) + " hours "
                                if difference.minutes == 1:
                                    minutes = "and 1 minute."
                                else:
                                    minutes = "and " + str(difference.minutes) + " minutes."
                            else:
                                days = "0 days "
                                hours = "0 hours "
                                minutes = "and 0 minutes."
                            ransomware_results.append("Ransomware activity was reported by end users within " + days + hours + minutes)
                        else:
                            ransomware_results.append("Ransomware activity was not reported by end users.")

            for r in reversed(ransomware_results):
                b = doc.add_paragraph()
                b.text = r
                b.style = "List Bullet 2"
                p_tag._p.addnext(b._p)

            if len(ransomware_results) > 0:
                customer_name = af.get_db_info(db, 'engagementmeta.fields.customer_long_name', '{STAKEHOLDER NAME}')
                results = doc.add_paragraph(f"Additionally, in coordination with the {customer_name} point-of-contact, the team identified the following detection and response results:", style="Normal")
                p_tag._p.addnext(results._p)

        if af.get_db_info(db, 'ransomwarescenarios', 'NA') is not None:
            vuln = 0
            total = 0
            if af.get_db_info(db, 'ransomwarescenarios.fields.vuln', 'NA') != '<not set: NA>':
                vuln = int(af.get_db_info(db, 'ransomwarescenarios.fields.vuln', 'NA'))
            if af.get_db_info(db, 'ransomwarescenarios.fields.total', 'NA') != '<not set: NA>':
                total = int(af.get_db_info(db, 'ransomwarescenarios.fields.total', 'NA'))

            customer_name = af.get_db_info(db, 'engagementmeta.fields.customer_long_name', '{STAKEHOLDER NAME}')
            sim = doc.add_paragraph(f"During ransomware simulation, the CISA team found that endpoints in the {customer_name} network are vulnerable to {vuln} out of the {total} ransomware scenarios tested. Details of the ransomware simulation can be found in the corresponding report provided with the scan reports and raw data.", style="Normal")
            p_tag._p.addnext(sim._p)

        para = doc.add_paragraph("Ransomware is malicious software that encrypts files on infected hosts. In order to decrypt affected files, users or organizations are instructed to pay a ransom. Infection occurs via manual or automated deployment after an attacker has compromised a host. The impact of mass infection can be severe including lost revenue, lost data, diminished reputation, and various costs associated with incident response and recovery.", style="Normal")
        p_tag._p.addnext(para._p)

        header = doc.add_paragraph("Ransomware Susceptibility", style="Heading 2")
        p_tag._p.addnext(header._p)

    xu.delete_paragraph(p_tag)


def insert_de_table(doc, db):
    """Function that fills out the data exfiltration section of the report if data exfiltration results exist.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{DATA EXFILTRATION RESULTS}")

    if p_tag is None:
        return

    if af.get_db_info(db, 'dataexfil', 'NA') is not None:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

        rec = doc.add_paragraph()
        r_run1 = rec.add_run("In addition to the recommended mitigations provided throughout the ")
        r_run2 = rec.add_run("Findings")
        r_run3 = rec.add_run(" section pertaining to protecting sensitive data, the CISA team recommends designing and implementing network perimeter controls to detect and prevent data exfiltration. Where possible, outbound traffic should be forced through authenticated proxy servers on the enterprise perimeter. These proxy servers can be configured to decrypt and inspect network traffic, flagging suspicious or anomalous characteristics. Additionally, perimeter controls should be configured with rules that allow or block traffic based on its destination.")
        rec.runs[1].font.bold = True
        rec.style = "Normal"
        p_tag._p.addnext(rec._p)

        caption = xu.insert_caption(doc, False, "Data Exfiltration Results")
        p_tag._p.addnext(caption._p)

        de_table = doc.add_table(1, 5)
        de_table.style = doc.styles['CISA Table']

        de_table.cell(0, 0).text = "Data Type"
        de_table.cell(0, 1).text = "Protocol"
        de_table.cell(0, 2).text = "Timestamp"
        de_table.cell(0, 3).text = "Detection"
        de_table.cell(0, 4).text = "Prevention"

        for cnt, data in enumerate(af.model_gen(db, "ptportal.dataexfil")):
            ele = data['fields']
            row = de_table.add_row()

            if ele['detection'] == "D":
                detection = "Detected"
            else:
                detection = "Not Detected"

            if ele['prevention'] == "B":
                prevention = "Blocked"
            else:
                prevention = "Not Blocked"

            if ele['date_time'] is not None:
                time = datetime.datetime.fromisoformat(ele['date_time'][:-1]).strftime("%m/%d/%Y %H:%M %p")
                timestamp = str(time) + " (UTC)"
            else:
                timestamp = ""

            de_table.cell(cnt + 1, 0).text = xu.xsafe(ele['datatype'])
            de_table.cell(cnt + 1, 1).text = xu.xsafe(ele['protocol'])
            de_table.cell(cnt + 1, 2).text = timestamp

            det_cell = de_table.cell(cnt + 1, 3).paragraphs[0]

            if detection == "Detected":
                det_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
            else:
                det_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)

            det_cell.add_run(detection)

            pre_cell = de_table.cell(cnt + 1, 4).paragraphs[0]

            if prevention == "Blocked":
                pre_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
            else:
                pre_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)

            pre_cell.add_run(prevention)

            de_table.cell(cnt + 1, 3).paragraphs[0].runs[0].font.bold = True
            de_table.cell(cnt + 1, 3).paragraphs[0].runs[1].font.bold = True
            de_table.cell(cnt + 1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            de_table.cell(cnt + 1, 4).paragraphs[0].runs[0].font.bold = True
            de_table.cell(cnt + 1, 4).paragraphs[0].runs[1].font.bold = True
            de_table.cell(cnt + 1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        de_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        de_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        de_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        de_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        de_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        de_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
        de_table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        de_table.cell(0, 3).paragraphs[0].runs[0].font.bold = True
        de_table.cell(0, 4).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        de_table.cell(0, 4).paragraphs[0].runs[0].font.bold = True

        xu.set_column_width(de_table.columns[0], 2.32)
        xu.set_column_width(de_table.columns[1], 0.74)
        xu.set_column_width(de_table.columns[2], 1.09)
        xu.set_column_width(de_table.columns[3], 1.17)
        xu.set_column_width(de_table.columns[4], 1.17)

        xu.move_table_after(de_table, p_tag)

        space = doc.add_paragraph("", style="Normal")
        p_tag._p.addnext(space._p)

        intro = doc.add_paragraph("The loss of sensitive business data often represents one of the highest risks to an enterprise. The CISA team conducted a data exfiltration simulation, in which fabricated data formatted to look like Personally Identifiable Information (PII) was transferred from an internal network to a CISA-controlled server outside of the network. The following results were observed when the CISA team attempted to exfiltrate data over common protocols:", style="Normal")
        p_tag._p.addnext(intro._p)

        header = doc.add_paragraph("Data Exfiltration", style="Heading 2")
        p_tag._p.addnext(header._p)

    xu.delete_paragraph(p_tag)


def insert_pt_table(doc, db):
    """Function that fills out the payload testing section of the report if payload testing results exist.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{PAYLOAD TESTING RESULTS}")

    if p_tag is None:
        return

    if af.get_db_info(db, 'payload', 'NA') is not None:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

        caption = xu.insert_caption(doc, False, "Payload Testing Results")
        p_tag._p.addnext(caption._p)

        pt_table = doc.add_table(1, 4)
        pt_table.style = doc.styles['CISA Table']

        pt_table.cell(0, 0).text = "Payload"
        pt_table.cell(0, 1).text = "C2 Protocol"
        pt_table.cell(0, 2).text = "Host Protection"
        pt_table.cell(0, 3).text = "Border Protection"

        for cnt, payload in enumerate(af.model_gen(db, "ptportal.payload")):
            ele = payload['fields']
            row = pt_table.add_row()

            if ele['host_protection'] == "B":
                host = "Blocked"
            else:
                host = "Not Blocked"

            if ele['border_protection'] == "B":
                border = "Blocked"
            else:
                border = "Not Blocked"

            pt_table.cell(cnt + 1, 0).text = xu.xsafe(ele['payload_description'])
            pt_table.cell(cnt + 1, 1).text = xu.xsafe(ele['c2_protocol'])

            host_cell = pt_table.cell(cnt + 1, 2).paragraphs[0]

            if host == "Blocked":
                host_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
            else:
                host_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)

            host_cell.add_run(host)

            bord_cell = pt_table.cell(cnt + 1, 3).paragraphs[0]

            if border == "Blocked":
                bord_cell.add_run("• ").font.color.rgb = RGBColor(131, 224, 142)
            else:
                bord_cell.add_run("• ").font.color.rgb = RGBColor(255, 116, 113)

            bord_cell.add_run(border)

            pt_table.cell(cnt + 1, 2).paragraphs[0].runs[0].font.bold = True
            pt_table.cell(cnt + 1, 2).paragraphs[0].runs[1].font.bold = True
            pt_table.cell(cnt + 1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_table.cell(cnt + 1, 3).paragraphs[0].runs[0].font.bold = True
            pt_table.cell(cnt + 1, 3).paragraphs[0].runs[1].font.bold = True
            pt_table.cell(cnt + 1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        pt_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        pt_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        pt_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        pt_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        pt_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        pt_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
        pt_table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        pt_table.cell(0, 3).paragraphs[0].runs[0].font.bold = True

        xu.set_column_width(pt_table.columns[0], 3.12)
        xu.set_column_width(pt_table.columns[1], 0.88)
        xu.set_column_width(pt_table.columns[2], 1.25)
        xu.set_column_width(pt_table.columns[3], 1.25)

        xu.move_table_after(pt_table, p_tag)

        space = doc.add_paragraph("", style="Normal")
        p_tag._p.addnext(space._p)

        third = doc.add_paragraph("For the purpose of this testing, border protection refers to the ability to make an outbound connection from the internal network to the CISA-controlled C2 server. If this connection can be made over a particular protocol, it is assumed that all payloads utilizing the same protocol would not be blocked at the border.", style="Normal")
        p_tag._p.addnext(third._p)

        second = doc.add_paragraph("While it is good to have strong filters and configurations in place to protect users from phishing attacks, it is important to note that a determined adversary could likely circumvent mail and browser filters if an established, trusted domain is used. For this reason, all recommendations outlined in this report are important considerations for defense-in-depth.", style="Normal")
        p_tag._p.addnext(second._p)

        customer_name = af.get_db_info(db, 'engagementmeta.fields.customer_long_name', '{STAKEHOLDER NAME}')
        exception = af.get_db_info(db, 'report.fields.exception', '{MANUAL EXCEPTION}')
        browser = af.get_db_info(db, 'report.fields.browser', '{INTERNET BROWSER}')
        intro = doc.add_paragraph(f"During payload testing with {customer_name}, a manual exception {exception} necessary for the test email to reach the target inbox. {browser} was used to download the payloads and was deemed a primary browser used by {customer_name} users.", style="Normal")
        p_tag._p.addnext(intro._p)

        header = doc.add_paragraph("Payload Testing Results", style="Heading 2")
        p_tag._p.addnext(header._p)

    xu.delete_paragraph(p_tag)


def insert_pc_table(doc, db):
    """Function that fills out the payload campaign section of the report if payload campaign data exists.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{PHISHING CAMPAIGN RESULTS}")

    if p_tag is None:
        return

    if af.get_db_info(db, 'campaign', 'NA') is not None:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

        for cnt, cam in reversed(list(enumerate(af.model_gen(db, "ptportal.campaign")))):
            ele = cam['fields']

            click_rate = round(float(ele['click_rate']) * 100, 2)
            click_time = str(ele['time_to_first_click']).split(" ")

            if len(click_time) > 1:
                if int(click_time[0]) == 1:
                    time_to_first_click = click_time[0] + " Day, " + click_time[1]
                else:
                    time_to_first_click = click_time[0] + " Days, " + click_time[1]
            else:
                time_to_first_click = str(ele['time_to_first_click'])

            if int(ele['length_of_campaign']) == 1:
                length = str(ele['length_of_campaign']) + " Day"
            else:
                length = str(ele['length_of_campaign']) + " Days"

            end_space = doc.add_paragraph("", style="Normal")
            p_tag._p.addnext(end_space._p)

            caption = xu.insert_caption(doc, False, "Phishing Campaign Results (" + str(cnt + 1) + ")")
            p_tag._p.addnext(caption._p)

            pc_table = doc.add_table(1, 2)
            pc_table.style = doc.styles['CISA Table']

            pc_table.cell(0, 0).merge(pc_table.cell(0, 1))
            pc_table.cell(0, 0).text = "Campaign #" + str(cnt + 1)

            row1 = pc_table.add_row()

            pc_table.cell(1, 0).text = "Emails Sent"
            pc_table.cell(1, 1).text = str(ele['emails_sent'])

            row2 = pc_table.add_row()

            pc_table.cell(2, 0).text = "Emails Successfully Delivered"
            pc_table.cell(2, 1).text = str(ele['emails_delivered'])

            row3 = pc_table.add_row()

            pc_table.cell(3, 0).text = "Click Rate"
            pc_table.cell(3, 1).text = str(click_rate) + "%"

            row4 = pc_table.add_row()

            pc_table.cell(4, 0).text = "Total Clicks"
            pc_table.cell(4, 1).text = str(ele['total_clicks'])

            row5 = pc_table.add_row()

            pc_table.cell(5, 0).text = "Unique Clicks"
            pc_table.cell(5, 1).text = str(ele['unique_clicks'])

            row6 = pc_table.add_row()

            pc_table.cell(6, 0).text = "Time to First Click (HH:MM:SS)"
            pc_table.cell(6, 1).text = time_to_first_click

            row7 = pc_table.add_row()

            pc_table.cell(7, 0).text = "Credentials Harvested"
            if str(ele['creds_harvested']) == "None":
                pc_table.cell(7, 1).text = "N/A"
            else:
                pc_table.cell(7, 1).text = str(ele['creds_harvested'])

            row8 = pc_table.add_row()

            pc_table.cell(8, 0).text = "Users Exploited"
            if str(ele['number_exploited']) == "None":
                pc_table.cell(8, 1).text = "N/A"
            else:
                pc_table.cell(8, 1).text = str(ele['number_exploited'])

            row9 = pc_table.add_row()

            pc_table.cell(9, 0).text = "Length of Campaign"
            pc_table.cell(9, 1).text = length

            pc_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            pc_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

            xu.set_column_width(pc_table.columns[0], 4)
            xu.set_column_width(pc_table.columns[1], 2.5)

            xu.move_table_after(pc_table, p_tag)

            space = doc.add_paragraph("", style="Normal")
            p_tag._p.addnext(space._p)

            description = xu.xsafe(ele['campaign_description'])
            campaign_description = doc.add_paragraph(description, style="Normal")
            p_tag._p.addnext(campaign_description._p)

        header = doc.add_paragraph("Phishing Campaign Results", style="Heading 2")
        p_tag._p.addnext(header._p)

    xu.delete_paragraph(p_tag)


def insert_osinf_tables(doc, db):
    """Function that fills out the OSINF section of the report if OSINF data exists.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """
    p_tag = xu.find_paragraph(doc, "{OSINF RESULTS}")

    if p_tag is None:
        return

    if af.get_db_info(db, 'breachedemail', 'NA') is not None:

        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

        caption_emails = xu.insert_caption(doc, False, "OSINF Breached Email Addresses")
        p_tag._p.addnext(caption_emails._p)

        be_table = doc.add_table(1, 2)
        be_table.style = doc.styles['CISA Table']

        be_table.cell(0, 0).text = "Email Address"
        be_table.cell(0, 1).text = "Breach Information"

        for cnt, email in enumerate(af.model_gen(db, "ptportal.breachedemail")):
            ele = email['fields']
            row = be_table.add_row()

            be_table.cell(cnt + 1, 0).text = xu.xsafe(ele['email_address'])
            be_table.cell(cnt + 1, 1).text = xu.xsafe(ele['breach_info'])

            be_table.cell(cnt + 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            be_table.cell(cnt + 1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


        be_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        be_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        be_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        be_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True

        xu.set_column_width(be_table.columns[0], 2.5)
        xu.set_column_width(be_table.columns[1], 4)

        xu.move_table_after(be_table, p_tag)

    if af.get_db_info(db, 'breachmetrics', 'NA') is not None and af.get_db_info(db, 'breachedemail', 'NA') is None:
        page_break = doc.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)
        p_tag._p.addnext(page_break._p)

    if af.get_db_info(db, 'breachmetrics', 'NA') is not None:
        if af.get_db_info(db, 'breachmetrics.fields.emails_identified', 'NA') != '<not set: NA>':
            emails_identified = int(af.get_db_info(db, 'breachmetrics.fields.emails_identified', 'NA'))
        else:
            emails_identified = 0

        if af.get_db_info(db, 'breachmetrics.fields.emails_identified_tp', 'NA') != '<not set: NA>':
            emails_identified_tp = int(af.get_db_info(db, 'breachmetrics.fields.emails_identified_tp', 'NA'))
        else:
            emails_identified_tp = 0

        if af.get_db_info(db, 'breachmetrics.fields.percentage_emails', 'NA') != '<not set: NA>':
            percentage_emails = round(float(af.get_db_info(db, 'breachmetrics.fields.percentage_emails', 'NA')) * 100, 2)
        else:
            percentage_emails = 0.0

        percentage = str(percentage_emails) + "%"

        if af.get_db_info(db, 'breachmetrics.fields.creds_identified', 'NA') != '<not set: NA>':
            creds_identified = int(af.get_db_info(db, 'breachmetrics.fields.creds_identified', 'NA'))
        else:
            creds_identified = 0

        if af.get_db_info(db, 'breachmetrics.fields.creds_identified_unique', 'NA') != '<not set: NA>':
            creds_identified_unique = int(af.get_db_info(db, 'breachmetrics.fields.creds_identified_unique', 'NA'))
        else:
            creds_identified_unique = 0

        if af.get_db_info(db, 'breachmetrics.fields.creds_validated', 'NA') != '<not set: NA>':
            creds_validated = int(af.get_db_info(db, 'breachmetrics.fields.creds_validated', 'NA'))
        else:
            creds_validated = 0

        end_space = doc.add_paragraph("", style="Normal")
        p_tag._p.addnext(end_space._p)

        caption_cred_metrics = xu.insert_caption(doc, False, "OSINF Breached Credential Metrics")
        p_tag._p.addnext(caption_cred_metrics._p)

        bm_creds_table = doc.add_table(1, 2)
        bm_creds_table.style = doc.styles['CISA Table']

        bm_creds_table.cell(0, 0).merge(bm_creds_table.cell(0, 1))
        bm_creds_table.cell(0, 0).text = "Breached Credential Metrics"

        row1 = bm_creds_table.add_row()

        bm_creds_table.cell(1, 0).text = "Credentials Identified"
        bm_creds_table.cell(1, 1).text = str(creds_identified)

        row2 = bm_creds_table.add_row()

        bm_creds_table.cell(2, 0).text = "Unique Users With Identified Credentials"
        bm_creds_table.cell(2, 1).text = str(creds_identified_unique)

        row3 = bm_creds_table.add_row()

        bm_creds_table.cell(3, 0).text = "Credentials Successfully Validated"
        bm_creds_table.cell(3, 1).text = str(creds_validated)

        bm_creds_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        bm_creds_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

        xu.set_column_width(bm_creds_table.columns[0], 4.75)
        xu.set_column_width(bm_creds_table.columns[1], 1.75)

        xu.move_table_after(bm_creds_table, p_tag)

        end_space = doc.add_paragraph("", style="Normal")
        p_tag._p.addnext(end_space._p)

        caption_email_metrics = xu.insert_caption(doc, False, "OSINF Breached Email Metrics")
        p_tag._p.addnext(caption_email_metrics._p)

        bm_email_table = doc.add_table(1, 2)
        bm_email_table.style = doc.styles['CISA Table']

        bm_email_table.cell(0, 0).merge(bm_email_table.cell(0, 1))
        bm_email_table.cell(0, 0).text = "Breached Email Metrics"

        row4 = bm_email_table.add_row()

        bm_email_table.cell(1, 0).text = "Emails Identified"
        bm_email_table.cell(1, 1).text = str(emails_identified)

        row5 = bm_email_table.add_row()

        bm_email_table.cell(2, 0).text = "Emails Identified in Third-Party Data Breaches"
        bm_email_table.cell(2, 1).text = str(emails_identified_tp)

        row6 = bm_email_table.add_row()

        bm_email_table.cell(3, 0).text = "Percentage of Emails Identified in Third-Party Data Breaches"
        bm_email_table.cell(3, 1).text = percentage

        bm_email_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        bm_email_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True

        xu.set_column_width(bm_email_table.columns[0], 4.75)
        xu.set_column_width(bm_email_table.columns[1], 1.75)

        xu.move_table_after(bm_email_table, p_tag)

        space = doc.add_paragraph("", style="Normal")
        p_tag._p.addnext(space._p)

        counts = {"emails_identified": emails_identified, "emails_identified_tp": emails_identified_tp, "creds_identified": creds_identified, "creds_identified_unique": creds_identified_unique, "creds_validated": creds_validated}
        words = {"emails_identified": "", "emails": "", "emails_identified_tp": "", "creds_identified": "", "creds_identified_unique": "", "creds_validated": ""}
        plural = {"emails_identified": "addresses", "emails": "These", "creds_identified_unique": "users", "creds_validated": "sets", "creds": "were"}

        if counts['emails_identified'] < 100:
            words['emails_identified'] = num.number_to_words(counts['emails_identified']) + " (" + str(counts['emails_identified']) + ")"
        else:
            words['emails_identified'] = str(counts['emails_identified'])

        if counts['emails_identified_tp'] < 100:
            words['emails_identified_tp'] = num.number_to_words(counts['emails_identified_tp']) + " (" + str(counts['emails_identified_tp']) + ")"
        else:
            words['emails_identified_tp'] = str(counts['emails_identified_tp'])

        if counts['creds_identified'] < 100:
            words['creds_identified'] = num.number_to_words(counts['creds_identified']) + " (" + str(counts['creds_identified']) + ")"
        else:
            words['creds_identified'] = str(counts['creds_identified'])

        if counts['creds_identified_unique'] < 100:
            words['creds_identified_unique'] = num.number_to_words(counts['creds_identified_unique']) + " (" + str(counts['creds_identified_unique']) + ")"
        else:
            words['creds_identified_unique'] = str(counts['creds_identified_unique'])

        if counts['creds_validated'] < 100:
            words['creds_validated'] = num.number_to_words(counts['creds_validated']) + " (" + str(counts['creds_validated']) + ")"
        else:
            words['creds_validated'] = str(counts['creds_validated'])

        if counts['emails_identified'] == 1:
            plural['emails_identified'] = "address"
            plural['emails'] = "This"

        if counts['creds_identified_unique'] == 1:
            plural['creds_identified_unique'] = "user"

        if counts['creds_validated'] == 1:
            plural['creds_validated'] = "set"
            plural['creds'] = "was"

        customer_name = af.get_db_info(db, 'engagementmeta.fields.customer_long_name', '{STAKEHOLDER NAME}')
        #emails_identified = af.get_db_info(db, 'breachmetrics.fields.emails_identified', '{EMAILS IDENTIFIED}')
        #emails_identified_tp = af.get_db_info(db, 'breachmetrics.fields.emails_identified_tp', '{EMAILS IDENTIFIED THIRD PARTY}')
        #percentage_emails = af.get_db_info(db, 'breachmetrics.fields.percentage_emails', '{PERCENTAGE EMAILS}')
        #creds_identified = af.get_db_info(db, 'breachmetrics.fields.creds_identified', '{CREDENTIALS IDENTIFIED}')
        #creds_identified_unique = af.get_db_info(db, 'breachmetrics.fields.creds_identified_unique', '{UNIQUE CREDENTIALS IDENTIFIED}')
        #creds_validated = af.get_db_info(db, 'breachmetrics.fields.creds_validated', '{CREDENTIALS VALIDATED}')

        third = doc.add_paragraph(f"Whenever credentials are discovered as a part of data breaches, CISA attempts to validate whether the credentials are still active and could be used as part of successful network exploitation. CISA was able to identify {words['creds_identified_unique']} unique {plural['creds_identified_unique']} with credentials found in publicly accessible breaches. Through testing, {words['creds_validated']} {plural['creds_validated']} of credentials {plural['creds']} successfully validated. Valid credentials can be used by threat actors to access {customer_name} resources.", style="Normal")
        p_tag._p.addnext(third._p)

        second = doc.add_paragraph(f"Once user email addresses have been identified, additional research is performed to determine if any have been exposed in publicly accessible breach information. CISA was able to identify that {words['emails_identified_tp']} of the discovered email addresses had been involved in breaches with publicly accessible data available. This indicates that {percentage} of identified users ({counts['emails_identified_tp']} of {counts['emails_identified']}) have been involved in previous data breaches.", style="Normal")
        p_tag._p.addnext(second._p)
        
        first = doc.add_paragraph(f"During the engagement, the CISA team performed reconnaissance of open-source information pertaining to the {customer_name} domain(s). With about one hour’s worth of effort, a threat actor would be able to identify {words['emails_identified']} unique user email {plural['emails_identified']} related to the {customer_name} domain. {plural['emails']} email {plural['emails_identified']} could be used in targeted social engineering attacks, such as phishing, or could be used to attempt brute force or password guessing attacks against authentication portals.", style="Normal")
        p_tag._p.addnext(first._p)

        header = doc.add_paragraph("Open-Source Information Gathering", style="Heading 2")
        p_tag._p.addnext(header._p)

    xu.delete_paragraph(p_tag)


def insert_pm_table(doc, db):
    """Function that fills out the port mapping appendix of the report.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{Table: Port Mapping}")

    if p_tag is None:
        return

    pm_table = doc.add_table(1, 4)
    pm_table.style = doc.styles['CISA Table']

    pm_table.cell(0, 0).text = "IP"
    pm_table.cell(0, 1).text = "Hostname"
    pm_table.cell(0, 2).text = "Ports"
    pm_table.cell(0, 3).text = "Services"

    for cnt, host in enumerate(af.model_gen(db, "ptportal.portmappinghost")):
        ele = host['fields']
        row = pm_table.add_row()

        ports = ele['ports'].split(", ")
        port_list = '\n'.join(ports)

        services = ele['services'].split(", ")
        service_list = '\n'.join(services)

        pm_table.cell(cnt + 1, 0).text = xu.xsafe(ele['ip'])
        pm_table.cell(cnt + 1, 1).text = xu.xsafe(ele['hostname'])
        pm_table.cell(cnt + 1, 2).text = port_list
        pm_table.cell(cnt + 1, 3).text = service_list

    pm_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    pm_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    pm_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    pm_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    pm_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    pm_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
    pm_table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    pm_table.cell(0, 3).paragraphs[0].runs[0].font.bold = True

    xu.set_column_width(pm_table.columns[0], 0.93)
    xu.set_column_width(pm_table.columns[1], 2.69)
    xu.set_column_width(pm_table.columns[2], 1.44)
    xu.set_column_width(pm_table.columns[3], 1.44)

    xu.move_table_after(pm_table, p_tag)

    xu.delete_paragraph(p_tag)


def insert_scope_table(doc, db):
    """Function that fills out the scope appendix of the report for FAST.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
    """

    p_tag = xu.find_paragraph(doc, "{Table: Scope}")

    if p_tag is None:
        return

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    p_tag._p.addnext(page_break._p)

    exc_table = doc.add_table(1, 2)
    exc_table.style = doc.styles['CISA Table']

    exc_table.cell(0, 0).text = "Service"
    exc_table.cell(0, 1).text = "IP Address / Hostname"

    exc_row = exc_table.add_row()
    exc_service = "Web Application and Penetration Test"
    exc_hosts = af.get_db_info(db, 'engagementmeta.fields.ext_excluded_scope', '{EXTERNAL EXCLUSIONS}')

    exc_table.cell(1, 0).text = exc_service
    exc_table.cell(1, 1).text = exc_hosts

    exc_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    exc_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    exc_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    exc_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True

    xu.set_column_width(exc_table.columns[0], 1.5)
    xu.set_column_width(exc_table.columns[1], 5)

    xu.move_table_after(exc_table, p_tag)

    para_break = doc.add_paragraph()
    p_tag._p.addnext(para_break._p)

    header = doc.add_paragraph("Exclusions", style="Heading 2")
    p_tag._p.addnext(header._p)

    para_break = doc.add_paragraph()
    p_tag._p.addnext(para_break._p)

    scope_table = doc.add_table(1, 3)
    scope_table.style = doc.styles['CISA Table']

    scope_table.cell(0, 0).text = "Service"
    scope_table.cell(0, 1).text = "IP Address / Hostname"
    scope_table.cell(0, 2).text = "Dates Performed"

    row = scope_table.add_row()
    service = "Web Application and Penetration Test"
    hosts = af.get_db_info(db, 'engagementmeta.fields.ext_scope', '{EXTERNAL SCOPE}')
    try:
        ext_start_date = af.get_db_info(db, 'engagementmeta.fields.ext_start_date', '{EXT START DATE}')
        start = datetime.datetime.strptime(ext_start_date, '%Y-%m-%d').strftime('%m/%d/%Y')
    except:
        start = "<NO DATE SET>"
    try:
        ext_end_date = af.get_db_info(db, 'engagementmeta.fields.ext_end_date', '{EXT END DATE}')
        end = datetime.datetime.strptime(ext_end_date, '%Y-%m-%d').strftime('%m/%d/%Y')
    except:
        end = "<NO DATE SET>"
    dates = start + " - " + end

    scope_table.cell(1, 0).text = service
    scope_table.cell(1, 1).text = hosts
    scope_table.cell(1, 2).text = dates

    scope_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    scope_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    scope_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    scope_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    scope_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    scope_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True

    xu.set_column_width(scope_table.columns[0], 1.5)
    xu.set_column_width(scope_table.columns[1], 3.8)
    xu.set_column_width(scope_table.columns[2], 1.2)

    xu.move_table_after(scope_table, p_tag)

    para_break = doc.add_paragraph()
    p_tag._p.addnext(para_break._p)

    xu.delete_paragraph(p_tag)


def generate_ptp_report(template, output, draft, json, media):
    """Generates a PTP report based on the provided json data.

    Args:
        template (string): Path to the docx template that will be used to generate the report.
        output (string): Name of the file that will be saved and returned.
        draft (boolean): Marks the docx document with a draft watermark.
        json (string): Path to the json file with the assessment data.
        media (string): Path to the media folder that contains the assessment screenshots.
    """
    if not os.path.exists(json):
        print("Invalid json file: ", json)
        sys.exit(1)
    if not os.path.exists(media):
        print("Invalid media path: ", media)
        sys.exit(1)
    if not os.path.exists(template):
        print("Invalid template file: ", template)
        sys.exit(1)

    # ---- Get data
    # gather meta, ndf, mam stats for charts, tables, etc.
    asmt_info = af.load_asmt_info(json)
    report_type = af.get_db_info(asmt_info, 'report.fields.report_type', '{REPORT TYPE}')

    if draft:
        af.set_draft(asmt_info)

    # ---- open the report template
    doc = docx.Document(template)

    # ---- replace paragraph tags
    for para in doc.paragraphs:
        for key in af.tag_db_map.keys():
            if key in para.text:
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if key == "{External Start Date}":
                            try:
                                ext_start_date = af.get_db_info(asmt_info, 'engagementmeta.fields.ext_start_date', '{EXT START DATE}')
                                replace_str = datetime.datetime.strptime(ext_start_date, '%Y-%m-%d').strftime('%B %-d, %Y')
                            except:
                                replace_str = "<NO DATE SET>"
                        elif key == "{External End Date}":
                            try:
                                ext_end_date = af.get_db_info(asmt_info, 'engagementmeta.fields.ext_end_date', '{EXT END DATE}')
                                replace_str = datetime.datetime.strptime(ext_end_date, '%Y-%m-%d').strftime('%B %-d, %Y')
                            except:
                                replace_str = "<NO DATE SET>"
                        elif key == "{Internal Start Date}":
                            try:
                                int_start_date = af.get_db_info(asmt_info, 'engagementmeta.fields.int_start_date', '{INT START DATE}')
                                replace_str = datetime.datetime.strptime(int_start_date, '%Y-%m-%d').strftime('%B %-d, %Y')
                            except:
                                replace_str = "<NO DATE SET>"
                        elif key == "{Internal End Date}":
                            try:
                                int_end_date = af.get_db_info(asmt_info, 'engagementmeta.fields.int_end_date', '{INT END DATE}')
                                replace_str = datetime.datetime.strptime(int_end_date, '%Y-%m-%d').strftime('%B %-d, %Y')
                            except:
                                replace_str = "<NO DATE SET>"
                        else:    
                            replace_str = af.get_db_info(asmt_info, af.tag_db_map[key], key)

                        if isinstance(replace_str, list):
                            replace_str = ', '.join(replace_str)

                        text = inline[i].text.replace(key, str(replace_str))
                        inline[i].text = text

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for key in af.tag_db_map.keys():
                    for para in cell.paragraphs:
                        if key in para.text:
                            replace_str = af.get_db_info(
                                asmt_info, af.tag_db_map[key], key
                            )
                            if isinstance(replace_str, list):
                                replace_str = ', '.join(replace_str)

                            text = para.text.replace(key, str(replace_str))
                            para.text = text

    insert_assessment_info(doc, asmt_info)
    
    if report_type == "RVA" or report_type == "RPT":
        # RVA 2.0 Template Insertions
        insert_report_summary(doc, asmt_info, media)
        insert_scope(doc, asmt_info)

        # ---- add risk score table
        insert_rs_table(doc, asmt_info, "{Table: Risk Score}", media)

        # ---- add cis_csc recommendations
        insert_cis_csc_table(doc, asmt_info, "{Table: CIS_CSC}")

    # ---- add findings and kevs
    insert_fs_table(doc, asmt_info, "{Table: Findings Summary}")
    insert_df_table(doc, asmt_info, "{Table: Detailed Findings}", media)
    #insert_kev_table(doc, asmt_info, "{Table: KEVs}")

    # ---- add additional service sections
    if report_type == "RVA":
        insert_ransomware(doc, asmt_info)
        insert_de_table(doc, asmt_info)
        insert_pt_table(doc, asmt_info)
        insert_pc_table(doc, asmt_info)
    if report_type == "RPT":
        insert_pt_table(doc, asmt_info)
        insert_osinf_tables(doc, asmt_info)
    if report_type == "FAST":
        insert_pc_table(doc, asmt_info)

    # ---- add attack paths
    insert_attack_paths(doc, asmt_info, media)

    # ---- add appendices
    insert_pm_table(doc, asmt_info)
    insert_narrative(doc, asmt_info, media)
    if report_type == "FAST":
        insert_scope_table(doc, asmt_info)
    if report_type == "RVA":
        insert_pass_analysis(doc, asmt_info)
    insert_acronyms(doc, asmt_info)

    tlp = af.get_db_info(asmt_info, 'engagementmeta.fields.traffic_light_protocol', '{TRAFFIC LIGHT PROTOCOL}')

    first_header = doc.sections[0].first_page_header
    first_footer = doc.sections[0].first_page_footer

    if tlp == "Red":
        label = " TLP:RED"
        color = RGBColor(255, 43, 43)
    elif tlp == "Amber":
        label = " TLP:AMBER"
        color = RGBColor(255, 192, 0)
    elif tlp == "Amber+Strict":
        label = " TLP:AMBER+STRICT"
        color = RGBColor(255, 192, 0)
    elif tlp == "Clear":
        label = " TLP:CLEAR"
        color = RGBColor(255, 255, 255)
    else:
        label = None

    if label is not None:
        for fheaderp in first_header.paragraphs:
            if "TLP" in fheaderp.text:
                p = fheaderp._p
            for run in fheaderp.runs:
                if "TLP" in run.text:
                    p.remove(run._r)
                    tlp_label = fheaderp.add_run(label)
                    tlp_label.font.color.rgb = color
                    tlp_label.font.name = "Franklin Gothic Medium"
                    tlp_label.font.size = docx.shared.Pt(14)
                    tlp_label.font.highlight_color = WD_COLOR_INDEX.BLACK
                    tlp_end = fheaderp.add_run(".")
                    tlp_end.font.color.rgb = RGBColor(0, 0, 0)
                    tlp_end.font.highlight_color = WD_COLOR_INDEX.BLACK

        for ffooterp in first_footer.paragraphs:
            if "TLP" in ffooterp.text:
                p = ffooterp._p
            for run in ffooterp.runs:
                if "TLP" in run.text:
                    p.remove(run._r)
                    tlp_label = ffooterp.add_run(label)
                    tlp_label.font.color.rgb = color
                    tlp_label.font.name = "Franklin Gothic Medium"
                    tlp_label.font.size = docx.shared.Pt(14)
                    tlp_label.font.highlight_color = WD_COLOR_INDEX.BLACK
                    tlp_end = ffooterp.add_run(".")
                    tlp_end.font.color.rgb = RGBColor(0, 0, 0)
                    tlp_end.font.highlight_color = WD_COLOR_INDEX.BLACK

        for section in doc.sections:
            header = section.header
            footer = section.footer
            for headerp in header.paragraphs:
                if "TLP" in headerp.text:
                    p = headerp._p
                    headerp.style = "TLP"
                    for run in headerp.runs:
                        p.remove(run._r)
                    tlp_label = headerp.add_run(label)
                    tlp_label.font.color.rgb = color
                    tlp_label.font.highlight_color = WD_COLOR_INDEX.BLACK
                    tlp_end = headerp.add_run(".")
                    tlp_end.font.color.rgb = RGBColor(0, 0, 0)
                    tlp_end.font.highlight_color = WD_COLOR_INDEX.BLACK

            for footerp in footer.paragraphs:
                if "TLP" in footerp.text:
                    p = footerp._p
                    footerp.style = "TLP"
                    for run in footerp.runs:
                        p.remove(run._r)
                    tlp_label = footerp.add_run(label)
                    tlp_label.font.color.rgb = color
                    tlp_label.font.highlight_color = WD_COLOR_INDEX.BLACK
                    tlp_end = footerp.add_run(".")
                    tlp_end.font.color.rgb = RGBColor(0, 0, 0)
                    tlp_end.font.highlight_color = WD_COLOR_INDEX.BLACK

    # ---- save the report
    doc.save(output)


def main():
    description = "Generate report"
    parser = argparse.ArgumentParser(description=description)

    parser.add_argument("TEMPLATE", help="Report template.")
    parser.add_argument(
        "-o",
        "--output_file",
        action="store",
        default="RVA_Report_" + tstamp + ".docx",
        help="Report file name",
    )
    parser.add_argument(
        "-d", "--draft", action="store_true", help="Label report as a draft"
    )
    parser.add_argument("-j", "--json_file", action="store", required=True)
    parser.add_argument(
        "-m",
        "--media_path",
        action="store",
        default="./",
        help="Location of screenshots, etc.",
    )
    args = parser.parse_args()

    generate_ptp_report(
        args.TEMPLATE, args.output_file, args.draft, args.json_file, args.media_path
    )


if __name__ == '__main__':
    main()