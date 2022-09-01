"""
The RVA report generator takes a document template and information
from the RVA pen testing django app database to generate an RVA report
in docx format
"""
# Risk & Vulnerability Reporting Engine

# Copyright 2022 Carnegie Mellon University.

# NO WARRANTY. THIS CARNEGIE MELLON UNIVERSITY AND SOFTWARE ENGINEERING INSTITUTE MATERIAL IS FURNISHED ON AN "AS-IS" BASIS. CARNEGIE MELLON UNIVERSITY MAKES NO WARRANTIES OF ANY KIND, EITHER EXPRESSED OR IMPLIED, AS TO ANY MATTER INCLUDING, BUT NOT LIMITED TO, WARRANTY OF FITNESS FOR PURPOSE OR MERCHANTABILITY, EXCLUSIVITY, OR RESULTS OBTAINED FROM USE OF THE MATERIAL. CARNEGIE MELLON UNIVERSITY DOES NOT MAKE ANY WARRANTY OF ANY KIND WITH RESPECT TO FREEDOM FROM PATENT, TRADEMARK, OR COPYRIGHT INFRINGEMENT.

# Released under a BSD (SEI)-style license, please see license.txt or contact permission@sei.cmu.edu for full terms.

# [DISTRIBUTION STATEMENT A] This material has been approved for public release and unlimited distribution.  Please see Copyright notice for non-US Government use and distribution.

# Carnegie Mellon® is registered in the U.S. Patent and Trademark Office by Carnegie Mellon University.

# This Software includes and/or makes use of Third-Party Software each subject to its own license.

# DM22-0744
import sys
import os.path
import argparse
import datetime

import report_gen.utilities.assessment_facts as af
import report_gen.utilities.xml_util as xu
import report_gen.utilities.rt_parser as rtp

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

try:
    import docx
except ImportError:
    print("Must have python docx library installed")
    sys.exit(1)


# ---- constants
tstamp = str(datetime.datetime.now().strftime("%Y%m%d_%H.%M.%S"))


def insert_risk_assessment(doc, db, mediapath):
    """Inserts the risk gauges into the report.

    Args:
        doc (docx document): The docx template being filled in.
        db (List of lists): Json format of assessment facts
        mediapath (String): Location of screenshots
    """

    # Build the findings lists for the verified exposures/risk assessments
    ite_exp = []
    dde_exp = []
    pe_exp = []
    re_exp = []
    for finding in af.model_gen(db, 'ptportal.uploadedfinding'):
        if finding.get("fields").get("insider_threat_exp") != 0.0:
            ite_exp.append(finding)
        if finding.get("fields").get("data_disclosure_exp") != 0.0:
            dde_exp.append(finding)
        if finding.get("fields").get("phishing_exp") != 0.0:
            pe_exp.append(finding)
        if finding.get("fields").get("ransomware_exp") != 0.0:
            re_exp.append(finding)

    table_tag = xu.find_paragraph(doc, '{Table: Verified Exposures}')
    # Makes it easier to then just loop through each entry associated with each section.
    risk_gauges = [
        ('Insider Threat', '/charts/itrisk.png', ite_exp),
        ('Phishing', '/charts/prisk.png', pe_exp),
        ('Ransomware', '/charts/rrisk.png', re_exp),
        ('Data Disclosure', '/charts/ddrisk.png', dde_exp),
    ]

    for gauge in risk_gauges[::-1]:
        p_tag = xu.find_paragraph(doc, "{" + gauge[0] + " Risk Gauge}")
        if p_tag is None:
            continue

        p = p_tag._p

        screen_p = doc.add_paragraph()
        r = screen_p.add_run()
        screen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r.add_picture(mediapath + gauge[1], width=Inches(5.24))
        p.addnext(screen_p._p)

        # Caption under the gauge.
        caption = xu.insert_caption(doc, gauge[0] + " Exposure")
        screen_p._p.addnext(caption._p)

        if table_tag is not None:
            # Setting up finding table and headers.
            finding_table = doc.add_table(len(gauge[2]) + 1, 3)
            finding_table.style = doc.styles['Report Default Table']
            finding_table.cell(0, 0).text = 'Finding'
            finding_table.cell(0, 1).text = 'Weight'
            finding_table.cell(0, 2).text = 'Affected Systems'

            # Get the Affected systems
            as_info = af.build_affected_systems_info(db)
            as_info = {k: xu.xsafe(v) for k, v in as_info.items()}

            # Enumerate through the findings list and populate the table.
            for num, f in enumerate(gauge[2]):
                finding_table.cell(num + 1, 0).text = f['fields'][
                    'uploaded_finding_name'
                ]

                if gauge[0] == "Insider Threat":
                    finding_table.cell(num + 1, 1).text = str(
                        f['fields']['insider_threat_exp']
                    )
                elif gauge[0] == "Phishing":
                    finding_table.cell(num + 1, 1).text = str(
                        f['fields']['phishing_exp']
                    )
                elif gauge[0] == "Ransomware":
                    finding_table.cell(num + 1, 1).text = str(
                        f['fields']['ransomware_exp']
                    )
                elif gauge[0] == "Data Disclosure":
                    finding_table.cell(num + 1, 1).text = str(
                        f['fields']['data_disclosure_exp']
                    )

                finding_table.cell(num + 1, 2).text = af.find_affected_systems(
                    as_info, f["fields"]["affected_systems"]
                )

            xu.move_table_after(finding_table, table_tag)
            table_tag._p.addnext(doc.add_paragraph(gauge[0], style='Heading 2')._p)

        xu.delete_paragraph(p_tag)

    if table_tag is not None:
        xu.delete_paragraph(table_tag)


def insert_report_summary(doc, db, media_path):
    """Function that fills out all the tags inside of the summary section of the report.

    Args:
        doc (docx object): The docx template.
        db (List of dictionaries): The JSON of the engagment data.
        media_path (String): Path to the media files
    """

    # Update all the charts in that section.
    xu.update_charts(doc, media_path, '{Table: CISA Findings}', 'fschart.png')
    xu.update_charts(doc, media_path, "{NIST 800-53 Controls}", "nistcontrol.png")
    xu.update_charts(doc, media_path, "{NIST CSF}", "nistcsf.png")

    # Creates the bulleted sections in the summary.
    bullet_lists = [
        ('{Cisa Results}', 'report.fields.cisa_results'),
        ('{Cisa Recommendations}', 'report.fields.cisa_recommendations'),
    ]
    for bl in bullet_lists:
        p_tag = xu.find_paragraph(doc, bl[0])
        if p_tag is not None:
            parser = rtp.RichTextParser(p_tag)
            parser.feed(af.get_db_info(db, bl[1], 'Results'))
            parser.emit_docx()
            parser.reset_parser()
            xu.delete_paragraph(p_tag)

    # Counts the number of findings of the severities.
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0, "Informational": 0}
    for finding in af.model_gen(db, 'ptportal.uploadedfinding'):
        counts[finding.get("fields").get("severity")] = (
            counts.get(finding.get("fields").get("severity"), 0) + 1
        )

    p_tag = xu.find_paragraph(doc, "{CISA Findings Summary}")
    if p_tag is not None:
        p_tag.text = f"CISA identified {counts['Critical']} critical-severity finding(s), {counts['High']} high-severity finding(s), {counts['Medium']} medium-severity finding(s), {counts['Low']} low-severity finding(s), and {counts['Informational']} informational-severity finding(s)."

    # Inserts all of the risk assessment section.
    insert_risk_assessment(doc, db, media_path)


def generate_ptp_summary(template, output, draft, json, media):
    """Generates a summary report for the current assessment.

    Args:
        template (string): Path to the docx template that will be used to generate the report.
        output (string): Name of the file that will be saved and returned.
        draft (boolean): Marks the docx document with a draft watermark.
        json (string): Path to the json file with the assessment data.
        media (string): Path to the media folder that contains the assessment screenshots.
    """
    # ---- verify files/directories exist
    if not os.path.exists(json):
        print("Invalid json file:", json)
        sys.exit(1)
    if not os.path.exists(media):
        print("Invalid media path:", media)
        # sys.exit(1)
    if not os.path.exists(template):
        print("Invalid template file:", template)
        sys.exit(1)

    # ---- Get data
    # gather meta, ndf, mam stats for charts, tables, etc.
    rva_info = af.load_rva_info(json)

    if draft:
        af.set_draft(rva_info)

    af.set_title(rva_info)

    # populate the nist information from the ndf data
    af.get_nist_control_data(rva_info)

    # ---- open the report template
    doc = docx.Document(template)

    # ---- update the document header
    xu.update_title(doc, rva_info)

    # ---- replace paragraph tags
    for para in doc.paragraphs:
        for key in af.tag_db_map.keys():
            if key in para.text:
                inline = para.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        replace_str = af.get_db_info(rva_info, af.tag_db_map[key], key)
                        if isinstance(replace_str, list):
                            replace_str = ', '.join(replace_str)

                        text = inline[i].text.replace(key, str(replace_str))
                        inline[i].text = text

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for key in af.tag_db_map.keys():
                    for para in cell.paragraphs:
                        if key in para.text:
                            replace_str = af.get_db_info(
                                rva_info, af.tag_db_map[key], key
                            )
                            if isinstance(replace_str, list):
                                replace_str = ', '.join(replace_str)

                            text = para.text.replace(key, str(replace_str))
                            para.text = text

    # RVA 2.0 Template Insertions
    insert_report_summary(doc, rva_info, media)

    # ---- save the report
    doc.save(output)


def main():
    description = "Generate RVA Report Summary of Results"
    parser = argparse.ArgumentParser(description=description)

    parser.add_argument("TEMPLATE", help="Report template.")
    parser.add_argument(
        "-o",
        "--output_file",
        action="store",
        default="RVA_Report_Summary_" + tstamp + ".docx",
        help="Report file name",
    )
    parser.add_argument(
        "-d", "--draft", action="store_true", help="Label report as a draft"
    )
    parser.add_argument("-j", "--json_file", action="store", required=True)
    parser.add_argument(
        "-m",
        "--media_path",
        action="store",
        default="./",
        help="Location of screenshots, etc.",
    )
    args = parser.parse_args()

    generate_ptp_summary(
        args.TEMPLATE, args.output_file, args.draft, args.json_file, args.media_path
    )


if __name__ == '__main__':
    main()
