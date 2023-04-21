"""
The RVA report generator takes a document template and information
from the RVA pen testing django app database to generate an RVA report
in docx format
"""
# Risk & Vulnerability Assessment Reporting Engine

# Copyright 2022 The Risk & Vulnerability Reporting Engine Contributors, All Rights Reserved.
# (see Contributors.txt for a full list of Contributors)

# SPDX-License-Identifier: BSD-3-Clause

# Please see additional acknowledgments (including references to third party source code, object code, documentation and other files) in the license.txt file or contact permission@sei.cmu.edu for full terms.

# Created, in part, with funding and support from the United States Government. (see Acknowledgments file).

# DM22-1011

import sys
import os.path
import argparse

from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT

import report_gen.utilities.assessment_facts as af
import report_gen.utilities.xml_util as xu
import report_gen.utilities.rt_parser as rtp

try:
    import docx
except ImportError:
    print("Must have python docx library installed")
    sys.exit(1)


def insert_kev_table(doc, db, kev_tag):
    """Generates and inserts the known exploited vulnerability table.

    Args:
        doc (docx document): The docx template being edited.
        db (list): Assessment data
        kev_tag (string): The string in the docx document that will be replaced with the table.
    """

    p_tag = xu.find_paragraph(doc, kev_tag)

    if p_tag is None:
        return

    kevs = []
    kev_ids = []

    for cnt, kev in enumerate(af.model_gen(db, "ptportal.kev")):
        ele = kev['fields']

        if ele['found']:
            found_kev = {"id": kev['pk'], "cve": ele['cve_id'], "vendor": ele['vendor_project'], "product": ele['product'], "vuln": ele['vulnerability_name'], "findings": []}
            kevs.append(found_kev)

    for cnt, finding in enumerate(af.model_gen(db, "ptportal.uploadedfinding")):
        ele = finding['fields']

        if ele['KEV']:        
            for item in kevs:
                if item['id'] in ele['KEV']:
                    item['findings'].append(ele['uploaded_finding_name'])

    kev_list = sorted(kevs, key=lambda d: d['cve']) 

    kev_table = doc.add_table(1, 5)
    kev_table.style = doc.styles['CISA Table']

    kev_table.cell(0, 0).text = "CVE"
    kev_table.cell(0, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 1).text = "Vendor"
    kev_table.cell(0, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 2).text = "Product"
    kev_table.cell(0, 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 3).text = "Vulnerability"
    kev_table.cell(0, 3).alignment = WD_ALIGN_PARAGRAPH.CENTER
    kev_table.cell(0, 4).text = "Finding(s)"
    kev_table.cell(0, 4).alignment = WD_ALIGN_PARAGRAPH.CENTER

    kev_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 3).paragraphs[0].runs[0].font.bold = True
    kev_table.cell(0, 4).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    kev_table.cell(0, 4).paragraphs[0].runs[0].font.bold = True

    for cnt, vuln in enumerate(kev_list):
        row = kev_table.add_row()

        url = "https://nvd.nist.gov/vuln/detail/" + vuln['cve']
        cve_cell = kev_table.cell(cnt + 1, 0)
        p = cve_cell.paragraphs[0]
        xu.add_link(p, 0, url, vuln['cve'])
        kev_table.cell(cnt + 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 1).text = vuln['vendor']
        kev_table.cell(cnt + 1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 2).text = vuln['product']
        kev_table.cell(cnt + 1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 3).text = vuln['vuln']
        kev_table.cell(cnt + 1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        kev_table.cell(cnt + 1, 4).text = ', '.join(vuln['findings'])
        kev_table.cell(cnt + 1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    xu.set_column_width(kev_table.columns[0], 0.6)
    xu.set_column_width(kev_table.columns[1], 0.72)
    xu.set_column_width(kev_table.columns[2], 0.81)
    xu.set_column_width(kev_table.columns[3], 2.19)
    xu.set_column_width(kev_table.columns[4], 2.19)

    xu.move_table_after(kev_table, p_tag)
    xu.delete_paragraph(p_tag)


def generate_kev_report(template, output, json, media):
    """Generates a PTP report based on the provided json data.

    Args:
        template (string): Path to the docx template that will be used to generate the report.
        output (string): Name of the file that will be saved and returned.
        draft (boolean): Marks the docx document with a draft watermark.
        json (string): Path to the json file with the assessment data.
        media (string): Path to the media folder that contains the assessment screenshots.
    """
    if not os.path.exists(json):
        print("Invalid json file: ", json)
        sys.exit(1)
    if not os.path.exists(media):
        print("Invalid media path: ", media)
        sys.exit(1)
    if not os.path.exists(template):
        print("Invalid template file: ", template)
        sys.exit(1)

    # ---- Get data
    # gather meta, ndf, mam stats for charts, tables, etc.
    rva_info = af.load_rva_info(json)

    # ---- open the report template and insert KEVs
    doc = docx.Document(template)
    insert_kev_table(doc, rva_info, "{Table: KEVs}")

    # ---- save the report
    doc.save(output)


def main():
    description = "Generate KEV report"
    parser = argparse.ArgumentParser(description=description)

    parser.add_argument("TEMPLATE", help="KEV template.")
    parser.add_argument(
        "-o",
        "--output_file",
        action="store",
        default="Known_Exploited_Vulnerabilities_Report.docx",
        help="KEV report file name",
    )
    parser.add_argument("-j", "--json_file", action="store", required=True)
    parser.add_argument(
        "-m",
        "--media_path",
        action="store",
        default="./",
        help="Location of screenshots, etc.",
    )
    args = parser.parse_args()

    generate_kev_report(
        args.TEMPLATE, args.output_file, args.json_file, args.media_path
    )


if __name__ == '__main__':
    main()
